import base64
import ast
import hashlib
import io
import html
import json
import math
import mimetypes
import os
import re
import sqlite3
import socket
import ipaddress
import threading
import time
import urllib.request
import urllib.error
import urllib.parse
import zipfile
import xml.etree.ElementTree as ET
import operator
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from html.parser import HTMLParser
from pathlib import Path
from pypdf import PdfReader
from docx import Document
from openpyxl import Workbook
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfgen import canvas

ROOT = Path(__file__).resolve().parent
DB = ROOT / "ai_workspace.db"
FILES = ROOT / "files"
LEGACY_ASSETS = ROOT / "legacy-assets"
HOST, PORT = "127.0.0.1", int(os.environ.get("AI_WORKSPACE_PORT", "8766"))
REQUIRE_AUTH = os.environ.get("AI_WORKSPACE_REQUIRE_AUTH", "false").strip().lower() in ("1", "true", "yes", "on")


class RateLimitError(ValueError):
    pass


def parse_model_ids(value):
    """Normalize a provider's optional manual model allowlist."""
    if isinstance(value, str):
        try:
            value = json.loads(value)
        except (TypeError, json.JSONDecodeError):
            value = re.split(r"[\n,]", value)
    if not isinstance(value, (list, tuple)):
        return []
    return list(dict.fromkeys(str(item).strip() for item in value if str(item).strip()))[:500]


def upstream_http_error(error, service="模型服务"):
    raw = error.read().decode("utf-8", errors="replace").strip()[:4000]
    detail = raw
    try:
        payload = json.loads(raw)
        error_value = payload.get("error", payload) if isinstance(payload, dict) else payload
        if isinstance(error_value, dict):
            detail = str(error_value.get("message") or error_value.get("detail") or error_value.get("status") or raw)
        elif error_value:
            detail = str(error_value)
    except (json.JSONDecodeError, TypeError):
        pass
    lowered = detail.lower()
    if "user location is not supported" in lowered or "location is not supported" in lowered:
        detail = (
            "Google Gemini 拒绝了当前服务器出口地区。生产服务器位于香港，Google Gemini API 对该出口地区不可用；"
            "这不是模型 ID 或 Key 格式问题。请改用支持当前地区的 Gemini 中转/OpenRouter，或把 AI 工作台部署到 Gemini 支持的地区。"
        )
    elif "valid api key" in lowered:
        detail = "API Key 无效，请重新粘贴 Google AI Studio 生成的 Gemini API Key"
    return ValueError(f"{service}返回 HTTP {error.code}: {detail or error.reason}")


def provider_api_base(base_url, provider_type="openai"):
    """Return the API root for a provider, accepting gateway roots as a convenience."""
    base = (base_url or "").rstrip("/")
    # Preset providers may use paths such as /v1beta/openai, /api/v3, or
    # /compatibility/v1. They are already API roots and must not receive a
    # second trailing /v1.
    api_root = r"(?:/v\d+(?:beta)?(?:/openai)?|/api/v\d+|/compatibility/v\d+)$"
    if provider_type != "ollama" and not re.search(api_root, base):
        return base + "/v1"
    return base


def infer_model_type(model_id):
    value = (model_id or "").lower()
    if any(token in value for token in ("gpt-image", "dall-e", "imagen", "flux", "stable-diffusion", "sdxl", "wanx")):
        return "image"
    if any(token in value for token in ("embedding", "embed-", "text-embedding")):
        return "embedding"
    if any(token in value for token in (
        "whisper", "transcribe", "transcription", "asr", "paraformer",
        "qwen-audio", "qwen-tts", "tts", "speech", "cosyvoice", "fun-asr",
        "sambert",
    )):
        return "audio"
    return "chat"


def provider_audio_presets(provider):
    """Return audio models that are part of the OpenAI-compatible API surface."""
    if not provider:
        return []
    provider_id = str(provider["provider_id"] if "provider_id" in provider.keys() else "").lower()
    base_url = str(provider["base_url"] if "base_url" in provider.keys() else "").lower()
    if provider_id == "openai" or "api.openai.com" in base_url:
        return ["gpt-4o-mini-transcribe", "gpt-4o-mini-tts"]
    return []


def audio_content_type(filename):
    suffix = Path(filename).suffix.lower()
    known = {
        ".flac": "audio/flac",
        ".mp3": "audio/mpeg",
        ".m4a": "audio/mp4",
        ".mp4": "audio/mp4",
        ".mpeg": "audio/mpeg",
        ".mpga": "audio/mpeg",
        ".ogg": "audio/ogg",
        ".wav": "audio/wav",
        ".webm": "audio/webm",
    }
    return known.get(suffix) or mimetypes.guess_type(filename)[0] or "application/octet-stream"


def extract_text(name, raw):
    suffix = Path(name).suffix.lower()
    if suffix == ".pdf":
        text = "\n\n".join((page.extract_text() or "").strip() for page in PdfReader(io.BytesIO(raw)).pages).strip()
    elif suffix == ".docx":
        with zipfile.ZipFile(io.BytesIO(raw)) as archive:
            root = ET.fromstring(archive.read("word/document.xml"))
        namespace = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        text = "\n".join("".join(node.text or "" for node in paragraph.iter(namespace + "t")).strip() for paragraph in root.iter(namespace + "p")).strip()
    elif suffix in (".txt", ".md", ".markdown", ".csv", ".json"):
        text = ""
        for encoding in ("utf-8-sig", "gb18030", "utf-16"):
            try:
                text = raw.decode(encoding).strip(); break
            except UnicodeDecodeError:
                continue
    elif suffix in (".png", ".jpg", ".jpeg", ".webp"):
        return "", "pending-ocr"
    else:
        raise ValueError("不支持该文件格式")
    if not text:
        raise ValueError("文件没有提取到可用文字")
    return text[:2_000_000], "ready"


def ocr_image(connection, name, raw, model_id=""):
    model = connection.execute("SELECT * FROM models WHERE id=? AND enabled=1", (model_id,)).fetchone() if model_id else None
    if not model:
        model = connection.execute("SELECT * FROM models WHERE enabled=1 AND (capabilities LIKE '%vision%' OR model_type='chat') ORDER BY is_default DESC,pinned DESC,name LIMIT 1").fetchone()
    if not model: raise ValueError("没有可用于 OCR 的视觉模型")
    provider = connection.execute("SELECT * FROM provider_connections WHERE id=? AND enabled=1", (model["connection_id"],)).fetchone() if model["connection_id"] else None
    if not provider: raise ValueError("OCR 模型没有可用连接")
    mime = mimetypes.guess_type(name)[0] or "image/png"
    image_url = f"data:{mime};base64,{base64.b64encode(raw).decode('ascii')}"
    payload = {"model": model["base_model"], "messages": [{"role": "user", "content": [{"type": "text", "text": "请完整提取图片中的全部文字，保持原有段落和表格顺序，只输出识别结果。"}, {"type": "image_url", "image_url": {"url": image_url}}]}], "stream": False}
    endpoint = provider_api_base(provider["base_url"], provider["provider_type"]) + "/chat/completions"
    headers = {"Content-Type": "application/json", "Authorization": "Bearer " + provider["api_key"], "User-Agent": "RuoShopAdmin/1.0"}
    request = urllib.request.Request(endpoint, data=json.dumps(payload).encode(), headers=headers, method="POST")
    with urllib.request.urlopen(request, timeout=180) as response: result = json.loads(response.read().decode("utf-8"))
    text = str(result.get("choices", [{}])[0].get("message", {}).get("content", "")).strip()
    if not text: raise ValueError("OCR 模型没有返回文字")
    return text[:2_000_000]


def generate_document(title, content, file_format):
    output = io.BytesIO(); file_format = file_format.lower()
    if file_format == "docx":
        document = Document(); document.add_heading(title, level=1)
        for paragraph in content.split("\n"): document.add_paragraph(paragraph)
        document.save(output); mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    elif file_format == "xlsx":
        workbook = Workbook(); sheet = workbook.active; sheet.title = title[:31] or "AI 输出"
        rows = [line.split("\t") if "\t" in line else [line] for line in content.splitlines()]
        for row in rows: sheet.append(row)
        sheet.freeze_panes = "A1"; workbook.save(output); mime = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    elif file_format == "pdf":
        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light")); pdf = canvas.Canvas(output, pagesize=A4)
        width, height = A4; text = pdf.beginText(48, height - 55); text.setFont("STSong-Light", 16); text.textLine(title); text.setFont("STSong-Light", 10)
        for paragraph in content.splitlines() or [""]:
            for start in range(0, len(paragraph), 52):
                if text.getY() < 48: pdf.drawText(text); pdf.showPage(); text = pdf.beginText(48, height - 48); text.setFont("STSong-Light", 10)
                text.textLine(paragraph[start:start + 52])
        pdf.drawText(text); pdf.save(); mime = "application/pdf"
    else: raise ValueError("仅支持 docx、xlsx 和 pdf")
    return output.getvalue(), mime


def multipart_file(fields, field_name, filename, raw, mime):
    boundary = "----RuoShopAdmin" + hashlib.sha256(raw[:1024]).hexdigest()[:16]
    parts = []
    for key, value in fields.items():
        parts.append(f"--{boundary}\r\nContent-Disposition: form-data; name=\"{key}\"\r\n\r\n{value}\r\n".encode())
    parts.append(f"--{boundary}\r\nContent-Disposition: form-data; name=\"{field_name}\"; filename=\"{filename}\"\r\nContent-Type: {mime}\r\n\r\n".encode() + raw + b"\r\n")
    parts.append(f"--{boundary}--\r\n".encode())
    return b"".join(parts), boundary


def web_search(query, limit=6):
    results = []
    try:
        params = urllib.parse.urlencode({"q": query, "format": "json", "no_html": "1", "skip_disambig": "1"})
        request = urllib.request.Request("https://api.duckduckgo.com/?" + params, headers={"User-Agent": "Mozilla/5.0"})
        with urllib.request.urlopen(request, timeout=12) as response: payload = json.loads(response.read().decode("utf-8"))
        if payload.get("AbstractText") and payload.get("AbstractURL"):
            results.append({"id": "web-0", "title": payload.get("Heading") or query, "url": payload["AbstractURL"], "content": payload["AbstractText"], "source": "web"})
        def collect(items):
            for item in items:
                if len(results) >= limit: return
                if item.get("Topics"): collect(item["Topics"])
                elif item.get("Text") and item.get("FirstURL"): results.append({"id": f"web-{len(results)}", "title": item["Text"].split(" - ", 1)[0][:160], "url": item["FirstURL"], "content": item["Text"], "source": "web"})
        collect(payload.get("RelatedTopics", []))
    except Exception:
        pass
    if not results:
        request = urllib.request.Request("https://www.bing.com/search?" + urllib.parse.urlencode({"q": query, "format": "rss"}), headers={"User-Agent": "Mozilla/5.0"})
        with urllib.request.urlopen(request, timeout=15) as response: root = ET.fromstring(response.read())
        for item in root.findall(".//item")[:limit]:
            title, url, description = item.findtext("title", ""), item.findtext("link", ""), item.findtext("description", "")
            clean = html.unescape(re.sub(r"<[^>]+>", " ", description)); clean = re.sub(r"\s+", " ", clean).strip()
            if title and url: results.append({"id": f"web-{len(results)}", "title": title, "url": url, "content": clean, "source": "web"})
    return results


class PublicPageParser(HTMLParser):
    def __init__(self):
        super().__init__(convert_charrefs=True)
        self.title_parts = []
        self.text_parts = []
        self._in_title = False
        self._skip_depth = 0

    def handle_starttag(self, tag, attrs):
        tag = tag.lower()
        if tag == "title":
            self._in_title = True
        if tag in {"script", "style", "noscript", "svg", "template"}:
            self._skip_depth += 1

    def handle_endtag(self, tag):
        tag = tag.lower()
        if tag == "title":
            self._in_title = False
        if tag in {"script", "style", "noscript", "svg", "template"} and self._skip_depth:
            self._skip_depth -= 1

    def handle_data(self, data):
        value = re.sub(r"\s+", " ", data).strip()
        if not value or self._skip_depth:
            return
        if self._in_title:
            self.title_parts.append(value)
        else:
            self.text_parts.append(value)


def read_public_web_page(url):
    parsed = urllib.parse.urlparse(str(url).strip())
    if parsed.scheme not in ("http", "https") or not parsed.hostname:
        raise ValueError("网页地址必须是 http 或 https")
    try:
        addresses = {item[4][0] for item in socket.getaddrinfo(parsed.hostname, parsed.port or (443 if parsed.scheme == "https" else 80))}
        if any(ipaddress.ip_address(address).is_private or ipaddress.ip_address(address).is_loopback or ipaddress.ip_address(address).is_reserved or ipaddress.ip_address(address).is_link_local for address in addresses):
            raise ValueError("网页读取禁止访问本机、内网或保留地址")
    except socket.gaierror as error:
        raise ValueError("网页地址无法解析") from error
    request = urllib.request.Request(parsed.geturl(), headers={"User-Agent": "Mozilla/5.0 RuoShopAdmin/1.0", "Accept": "text/html,application/xhtml+xml,text/plain;q=0.8"})
    with urllib.request.urlopen(request, timeout=20) as response:
        raw = response.read(2_000_000)
        charset = response.headers.get_content_charset() or "utf-8"
    source = raw.decode(charset, errors="replace")
    parser = PublicPageParser()
    parser.feed(source)
    content = re.sub(r"\s+", " ", " ".join(parser.text_parts)).strip()[:40_000]
    if not content:
        raise ValueError("网页没有可读取的正文")
    title = re.sub(r"\s+", " ", " ".join(parser.title_parts)).strip()[:200] or parsed.hostname
    return {"id": "web-page-" + hashlib.sha256(parsed.geturl().encode()).hexdigest()[:16], "title": title, "url": parsed.geturl(), "content": content, "source": "web-page"}


def message_text(content):
    """Extract plain text from the string or multimodal content we persist."""
    if isinstance(content, str):
        return content.strip()
    if isinstance(content, list):
        parts = []
        for item in content:
            if isinstance(item, str):
                parts.append(item)
            elif isinstance(item, dict) and isinstance(item.get("text"), str):
                parts.append(item["text"])
        return "\n".join(part.strip() for part in parts if part.strip()).strip()
    if isinstance(content, dict) and isinstance(content.get("text"), str):
        return content["text"].strip()
    return ""


def normalize_chat_messages(raw_messages, limit=24):
    """Convert saved/client messages into provider-compatible chat messages."""
    if not isinstance(raw_messages, list):
        return []
    result = []
    for item in raw_messages:
        if not isinstance(item, dict):
            continue
        if item.get("status") in ("failed", "cancelled"):
            continue
        role = str(item.get("role", "")).strip().lower()
        if role not in ("user", "assistant", "system"):
            continue
        text = message_text(item.get("content", ""))
        image_urls = item.get("image_urls", item.get("imageUrls", []))
        if not isinstance(image_urls, list):
            image_urls = []
        image_urls = [str(url) for url in image_urls if str(url).startswith(("data:image/", "https://", "http://"))][:4]
        if not text and not image_urls:
            continue
        if image_urls and role == "user":
            content = [{"type": "text", "text": text or "请分析这些图片"}]
            content.extend({"type": "image_url", "image_url": {"url": url}} for url in image_urls)
        else:
            content = text
        result.append({"role": role, "content": content})
    return result[-limit:]


def conversation_messages(connection, user_id, chat_id, fallback_messages):
    """Load an owned conversation; use client history only for unsaved sessions."""
    stored = []
    if chat_id:
        row = connection.execute("SELECT messages FROM chats WHERE id=? AND user_id=?", (chat_id, user_id)).fetchone()
        if row:
            try:
                stored = normalize_chat_messages(json.loads(row["messages"] or "[]"))
            except (TypeError, json.JSONDecodeError):
                stored = []
    return stored or normalize_chat_messages(fallback_messages)


def file_documents(connection, file_ids):
    """Return imported file text as model context for an explicit chat attachment."""
    if not isinstance(file_ids, list):
        return []
    ids = list(dict.fromkeys(str(item).strip() for item in file_ids if str(item).strip()))[:8]
    if not ids:
        return []
    placeholders = ",".join("?" for _ in ids)
    rows = connection.execute(f"SELECT id,name,content,status FROM files WHERE id IN ({placeholders})", ids).fetchall()
    return [
        {"id": row["id"], "title": row["name"], "content": row["content"][:20000], "status": row["status"], "source": "attachment"}
        for row in rows
        if row["content"] and row["status"] == "ready"
    ]


def db():
    connection = sqlite3.connect(DB)
    connection.row_factory = sqlite3.Row
    connection.executescript("""
      CREATE TABLE IF NOT EXISTS settings (key TEXT PRIMARY KEY, value TEXT NOT NULL);
      CREATE TABLE IF NOT EXISTS models (id TEXT PRIMARY KEY, name TEXT NOT NULL, base_model TEXT NOT NULL, description TEXT DEFAULT '', system_prompt TEXT DEFAULT '', capabilities TEXT DEFAULT '[]', updated_at INTEGER NOT NULL);
      CREATE TABLE IF NOT EXISTS provider_connections (id TEXT PRIMARY KEY, name TEXT NOT NULL, base_url TEXT NOT NULL, api_key TEXT NOT NULL DEFAULT '', enabled INTEGER NOT NULL DEFAULT 1, provider_type TEXT NOT NULL DEFAULT 'openai', purpose TEXT NOT NULL DEFAULT 'general', model_ids TEXT DEFAULT '[]', updated_at INTEGER NOT NULL);
      CREATE TABLE IF NOT EXISTS knowledge (id TEXT PRIMARY KEY, name TEXT NOT NULL, description TEXT DEFAULT '', created_at INTEGER NOT NULL);
      CREATE TABLE IF NOT EXISTS files (id TEXT PRIMARY KEY, knowledge_id TEXT, name TEXT NOT NULL, content TEXT DEFAULT '', path TEXT DEFAULT '', status TEXT DEFAULT 'ready', created_at INTEGER NOT NULL);
      CREATE TABLE IF NOT EXISTS file_chunks (id TEXT PRIMARY KEY, file_id TEXT NOT NULL, chunk_index INTEGER NOT NULL, content TEXT NOT NULL, embedding TEXT DEFAULT '');
      CREATE VIRTUAL TABLE IF NOT EXISTS file_chunks_fts USING fts5(chunk_id UNINDEXED, file_id UNINDEXED, title, content, tokenize='trigram');
      CREATE TABLE IF NOT EXISTS prompts (id TEXT PRIMARY KEY, command TEXT UNIQUE NOT NULL, title TEXT NOT NULL, content TEXT NOT NULL, updated_at INTEGER NOT NULL);
      CREATE TABLE IF NOT EXISTS notes (id TEXT PRIMARY KEY, title TEXT NOT NULL, content TEXT NOT NULL, updated_at INTEGER NOT NULL);
      CREATE TABLE IF NOT EXISTS shared_chats (id TEXT PRIMARY KEY, owner_id TEXT NOT NULL, title TEXT NOT NULL, messages TEXT NOT NULL, created_at INTEGER NOT NULL);
      CREATE TABLE IF NOT EXISTS skills (id TEXT PRIMARY KEY, name TEXT NOT NULL, description TEXT DEFAULT '', content TEXT NOT NULL, updated_at INTEGER NOT NULL);
      CREATE TABLE IF NOT EXISTS tools (id TEXT PRIMARY KEY, name TEXT NOT NULL, description TEXT DEFAULT '', kind TEXT NOT NULL, enabled INTEGER NOT NULL DEFAULT 1, updated_at INTEGER NOT NULL);
      CREATE TABLE IF NOT EXISTS chats (id TEXT PRIMARY KEY, user_id TEXT NOT NULL, title TEXT NOT NULL, messages TEXT NOT NULL DEFAULT '[]', folder TEXT DEFAULT '', archived INTEGER DEFAULT 0, created_at INTEGER NOT NULL, updated_at INTEGER NOT NULL);
      CREATE TABLE IF NOT EXISTS ai_usage (id TEXT PRIMARY KEY, user_id TEXT NOT NULL, model_id TEXT DEFAULT '', operation TEXT NOT NULL, input_tokens INTEGER DEFAULT 0, output_tokens INTEGER DEFAULT 0, cost REAL DEFAULT 0, latency_ms INTEGER DEFAULT 0, status TEXT NOT NULL, detail TEXT DEFAULT '', created_at INTEGER NOT NULL);
      CREATE TABLE IF NOT EXISTS ai_memories (id TEXT PRIMARY KEY, user_id TEXT NOT NULL, content TEXT NOT NULL, source_chat_id TEXT DEFAULT '', enabled INTEGER DEFAULT 1, created_at INTEGER NOT NULL, updated_at INTEGER NOT NULL);
      CREATE TABLE IF NOT EXISTS ai_jobs (id TEXT PRIMARY KEY, user_id TEXT NOT NULL, kind TEXT NOT NULL, status TEXT NOT NULL, input TEXT DEFAULT '{}', output TEXT DEFAULT '{}', error TEXT DEFAULT '', created_at INTEGER NOT NULL, updated_at INTEGER NOT NULL);
      CREATE TABLE IF NOT EXISTS ai_workflows (id TEXT PRIMARY KEY, owner_id TEXT NOT NULL, name TEXT NOT NULL, description TEXT DEFAULT '', steps TEXT DEFAULT '[]', enabled INTEGER DEFAULT 1, created_at INTEGER NOT NULL, updated_at INTEGER NOT NULL);
      CREATE TABLE IF NOT EXISTS ai_rate_limits (user_id TEXT NOT NULL, bucket INTEGER NOT NULL, request_count INTEGER DEFAULT 0, PRIMARY KEY(user_id,bucket));
    """)
    model_columns = {row["name"] for row in connection.execute("PRAGMA table_info(models)")}
    provider_columns = {row["name"] for row in connection.execute("PRAGMA table_info(provider_connections)")}
    if "provider_type" not in provider_columns: connection.execute("ALTER TABLE provider_connections ADD COLUMN provider_type TEXT NOT NULL DEFAULT 'openai'")
    if "purpose" not in provider_columns: connection.execute("ALTER TABLE provider_connections ADD COLUMN purpose TEXT NOT NULL DEFAULT 'general'")
    if "provider_id" not in provider_columns: connection.execute("ALTER TABLE provider_connections ADD COLUMN provider_id TEXT NOT NULL DEFAULT 'custom'")
    if "model_ids" not in provider_columns: connection.execute("ALTER TABLE provider_connections ADD COLUMN model_ids TEXT DEFAULT '[]'")
    provider_patterns = {
        "openai": ("api.openai.com",), "google": ("generativelanguage.googleapis.com",), "mistral": ("api.mistral.ai",),
        "groq": ("api.groq.com",), "xai": ("api.x.ai",), "openrouter": ("openrouter.ai",), "cohere": ("api.cohere.ai",),
        "deepseek": ("api.deepseek.com",), "qwen": ("dashscope.aliyuncs.com",), "zhipu": ("open.bigmodel.cn",),
        "moonshot": ("api.moonshot.cn",), "baichuan": ("api.baichuan-ai.com",), "yi": ("api.lingyiwanwu.com",),
        "doubao": ("volces.com",), "siliconflow": ("siliconflow.cn",), "minimax": ("minimax",), "stepfun": ("stepfun.com",),
        "ollama": ("127.0.0.1:11434", "localhost:11434"), "lmstudio": ("127.0.0.1:1234", "localhost:1234"),
        "localai": ("127.0.0.1:8080", "localhost:8080"), "pipeline": ("127.0.0.1:9099", "localhost:9099")
    }
    for provider_id, patterns in provider_patterns.items():
        for pattern in patterns:
            connection.execute("UPDATE provider_connections SET provider_id=? WHERE provider_id='custom' AND lower(base_url) LIKE ?", (provider_id, f"%{pattern}%"))
    for name, definition in {"temperature": "REAL DEFAULT 0.7", "top_p": "REAL DEFAULT 1", "max_tokens": "INTEGER DEFAULT 2048", "knowledge_id": "TEXT DEFAULT ''", "skill_ids": "TEXT DEFAULT '[]'", "tool_ids": "TEXT DEFAULT '[]'", "connection_id": "TEXT DEFAULT ''", "provider_id": "TEXT DEFAULT 'custom'", "enabled": "INTEGER DEFAULT 1", "hidden": "INTEGER DEFAULT 0", "pinned": "INTEGER DEFAULT 0", "is_default": "INTEGER DEFAULT 0", "tags": "TEXT DEFAULT '[]'", "sort_order": "INTEGER DEFAULT 0", "access": "TEXT DEFAULT 'private'", "access_grants": "TEXT DEFAULT '[]'", "filters": "TEXT DEFAULT '[]'", "actions": "TEXT DEFAULT '[]'", "owner_id": "TEXT DEFAULT 'local'", "model_type": "TEXT DEFAULT 'chat'", "input_price": "REAL DEFAULT 0", "output_price": "REAL DEFAULT 0"}.items():
        if name not in model_columns: connection.execute(f"ALTER TABLE models ADD COLUMN {name} {definition}")
    connection.execute("UPDATE models SET provider_id=COALESCE((SELECT provider_id FROM provider_connections WHERE provider_connections.id=models.connection_id),'custom') WHERE provider_id='custom' AND connection_id<>''")
    chunk_columns = {row["name"] for row in connection.execute("PRAGMA table_info(file_chunks)")}
    if "embedding" not in chunk_columns: connection.execute("ALTER TABLE file_chunks ADD COLUMN embedding TEXT DEFAULT ''")
    chat_columns = {row["name"] for row in connection.execute("PRAGMA table_info(chats)")}
    if "folder" not in chat_columns: connection.execute("ALTER TABLE chats ADD COLUMN folder TEXT DEFAULT ''")
    if "archived" not in chat_columns: connection.execute("ALTER TABLE chats ADD COLUMN archived INTEGER DEFAULT 0")
    if "model_id" not in chat_columns: connection.execute("ALTER TABLE chats ADD COLUMN model_id TEXT DEFAULT ''")
    if "favorite" not in chat_columns: connection.execute("ALTER TABLE chats ADD COLUMN favorite INTEGER DEFAULT 0")
    if "parent_chat_id" not in chat_columns: connection.execute("ALTER TABLE chats ADD COLUMN parent_chat_id TEXT DEFAULT ''")
    file_columns = {row["name"] for row in connection.execute("PRAGMA table_info(files)")}
    if "source" not in file_columns: connection.execute("ALTER TABLE files ADD COLUMN source TEXT DEFAULT ''")
    if "metadata" not in file_columns: connection.execute("ALTER TABLE files ADD COLUMN metadata TEXT DEFAULT '{}'")
    tool_columns = {row["name"] for row in connection.execute("PRAGMA table_info(tools)")}
    if "config" not in tool_columns: connection.execute("ALTER TABLE tools ADD COLUMN config TEXT DEFAULT '{}'")
    note_columns = {row["name"] for row in connection.execute("PRAGMA table_info(notes)")}
    if "owner_id" not in note_columns: connection.execute("ALTER TABLE notes ADD COLUMN owner_id TEXT NOT NULL DEFAULT 'local'")
    share_columns = {row["name"] for row in connection.execute("PRAGMA table_info(shared_chats)")}
    if "expires_at" not in share_columns: connection.execute("ALTER TABLE shared_chats ADD COLUMN expires_at INTEGER DEFAULT 0")
    if "revoked" not in share_columns: connection.execute("ALTER TABLE shared_chats ADD COLUMN revoked INTEGER DEFAULT 0")
    for row in connection.execute("SELECT id,base_model,model_type FROM models"):
        inferred = infer_model_type(row["base_model"])
        if not row["model_type"] or (row["model_type"] == "chat" and inferred != "chat"): connection.execute("UPDATE models SET model_type=? WHERE id=?", (inferred, row["id"]))
    now = int(time.time())
    if not connection.execute("SELECT 1 FROM provider_connections LIMIT 1").fetchone():
        values = {row["key"]: row["value"] for row in connection.execute("SELECT key,value FROM settings WHERE key IN ('base_url','api_key')")}
        if values.get("base_url"):
            connection.execute("INSERT INTO provider_connections(id,name,base_url,api_key,enabled,updated_at) VALUES(?,?,?,?,1,?)", ("legacy-default", "OpenAI 接口", values["base_url"], values.get("api_key", ""), now))
    connection.execute("INSERT OR IGNORE INTO tools(id,name,description,kind,updated_at) VALUES('builtin-calculator','计算器','安全计算加减乘除和括号表达式','calculator',?)", (now,))
    connection.execute("INSERT OR IGNORE INTO tools(id,name,description,kind,updated_at) VALUES('builtin-time','当前时间','读取服务器当前日期和时间','current_time',?)", (now,))
    connection.commit()
    return connection


def record_usage(connection, user_id, model_id, operation, started_at, status="ok", detail="", usage=None, cost=0):
    usage = usage or {}; now = int(time.time()); item_id = "usage-" + hashlib.sha256(f"{user_id}:{operation}:{time.time_ns()}".encode()).hexdigest()[:24]
    connection.execute("INSERT INTO ai_usage(id,user_id,model_id,operation,input_tokens,output_tokens,cost,latency_ms,status,detail,created_at) VALUES(?,?,?,?,?,?,?,?,?,?,?)", (item_id, user_id, model_id or "", operation, int(usage.get("prompt_tokens", usage.get("input_tokens", 0)) or 0), int(usage.get("completion_tokens", usage.get("output_tokens", 0)) or 0), max(0, float(cost or 0)), int((time.perf_counter() - started_at) * 1000), status, str(detail)[:1000], now))
    connection.commit()


def usage_cost(model, usage):
    if not model or not usage: return 0
    input_tokens = int(usage.get("prompt_tokens", usage.get("input_tokens", 0)) or 0); output_tokens = int(usage.get("completion_tokens", usage.get("output_tokens", 0)) or 0)
    return input_tokens / 1_000_000 * float(model["input_price"] or 0) + output_tokens / 1_000_000 * float(model["output_price"] or 0)


def upstream_event_value(value):
    """Extract text from common OpenAI-compatible, Responses, Gemini, and Ollama events."""
    if isinstance(value, str):
        return value
    if isinstance(value, list):
        return "".join(upstream_event_value(item) for item in value)
    if not isinstance(value, dict):
        return ""
    for key in ("delta", "text", "content", "output_text"):
        candidate = value.get(key)
        if isinstance(candidate, str):
            return candidate
        if isinstance(candidate, (dict, list)):
            text = upstream_event_value(candidate)
            if text:
                return text
    for key in ("message", "response", "output", "candidates", "choices", "parts"):
        text = upstream_event_value(value.get(key))
        if text:
            return text
    return ""


def upstream_event_usage(value):
    if not isinstance(value, dict):
        return {}
    usage = value.get("usage")
    if isinstance(usage, dict):
        return usage
    for key in ("response", "data", "output"):
        usage = upstream_event_usage(value.get(key))
        if usage:
            return usage
    return {}


def calculate_expression(question):
    matches = [item.strip() for item in re.findall(r"[\d.()]+(?:\s*[-+*/%]\s*[\d.()]+)+", question)]
    if not matches: return "未发现可计算的算式"
    allowed = {ast.Add: operator.add, ast.Sub: operator.sub, ast.Mult: operator.mul, ast.Div: operator.truediv, ast.Mod: operator.mod, ast.Pow: operator.pow, ast.USub: operator.neg, ast.UAdd: operator.pos}
    def evaluate(node):
        if isinstance(node, ast.Expression): return evaluate(node.body)
        if isinstance(node, ast.Constant) and isinstance(node.value, (int, float)): return node.value
        if isinstance(node, ast.BinOp) and type(node.op) in allowed: return allowed[type(node.op)](evaluate(node.left), evaluate(node.right))
        if isinstance(node, ast.UnaryOp) and type(node.op) in allowed: return allowed[type(node.op)](evaluate(node.operand))
        raise ValueError("算式包含不支持的内容")
    results = []
    for expression in matches[:5]: results.append(f"{expression} = {evaluate(ast.parse(expression, mode='eval'))}")
    return "\n".join(results)


def run_tools(connection, tool_ids, question):
    if not tool_ids: return ""
    placeholders = ",".join("?" for _ in tool_ids); rows = connection.execute(f"SELECT name,description,kind FROM tools WHERE enabled=1 AND id IN ({placeholders})", tool_ids).fetchall(); outputs = []
    for row in rows:
        if row["kind"] == "calculator": value = calculate_expression(question)
        elif row["kind"] == "current_time": value = time.strftime("%Y-%m-%d %H:%M:%S %z")
        elif row["kind"] == "custom": value = row["description"] or "按该工具名称所描述的职责处理用户请求。"
        else: continue
        outputs.append(f"## Tool: {row['name']}\n{value}")
    return "\n\n".join(outputs)


def tool_specs(connection, tool_ids):
    if not tool_ids: return [], {}
    placeholders = ",".join("?" for _ in tool_ids)
    rows = connection.execute(f"SELECT id,name,description,kind,config FROM tools WHERE enabled=1 AND id IN ({placeholders})", tool_ids).fetchall()
    specs, by_name = [], {}
    for index, row in enumerate(rows):
        function_name = re.sub(r"[^a-zA-Z0-9_]", "_", row["kind"] or row["id"]).strip("_") or f"tool_{index}"
        if function_name in by_name: function_name += f"_{index}"
        properties = {"expression": {"type": "string", "description": "要计算的算术表达式"}} if row["kind"] == "calculator" else {}
        required = ["expression"] if row["kind"] == "calculator" else []
        specs.append({"type": "function", "function": {"name": function_name, "description": row["description"] or row["name"], "parameters": {"type": "object", "properties": properties, "required": required, "additionalProperties": False}}})
        by_name[function_name] = dict(row)
    return specs, by_name


def execute_tool(row, arguments):
    if row["kind"] == "calculator":
        expression = str(arguments.get("expression", "")).strip()
        if not expression: return "缺少 expression 参数"
        return calculate_expression(expression) or "表达式无法计算"
    if row["kind"] == "current_time": return time.strftime("%Y-%m-%d %H:%M:%S %z")
    if row["kind"] == "http":
        config = json.loads(row.get("config") or "{}") if isinstance(row, dict) else json.loads(row["config"] or "{}")
        url = str(config.get("url", "")).strip(); parsed = urllib.parse.urlparse(url)
        if parsed.scheme not in ("http", "https") or not parsed.hostname: return "HTTP 工具地址无效"
        addresses = {item[4][0] for item in socket.getaddrinfo(parsed.hostname, parsed.port or (443 if parsed.scheme == "https" else 80))}
        if any(ipaddress.ip_address(address).is_private or ipaddress.ip_address(address).is_loopback or ipaddress.ip_address(address).is_reserved or ipaddress.ip_address(address).is_link_local for address in addresses): return "HTTP 工具禁止访问本机、内网或保留地址"
        method = str(config.get("method", "POST")).upper(); headers = {str(k): str(v) for k, v in dict(config.get("headers") or {}).items()}; headers.setdefault("Content-Type", "application/json")
        body = None if method == "GET" else json.dumps(arguments, ensure_ascii=False).encode("utf-8")
        if method == "GET" and arguments: url += ("&" if "?" in url else "?") + urllib.parse.urlencode(arguments)
        with urllib.request.urlopen(urllib.request.Request(url, data=body, headers=headers, method=method), timeout=min(30, max(1, int(config.get("timeout", 15))))) as response: return response.read(200000).decode("utf-8", errors="replace")
    if row["kind"] == "custom": return row["description"] or "工具已执行"
    return "不支持的工具类型"


def complete_with_tools(endpoint, headers, payload, specs, tool_map, max_rounds=4):
    messages = list(payload["messages"])
    for _ in range(max_rounds):
        request_payload = {**payload, "messages": messages, "stream": False, "tools": specs, "tool_choice": "auto"}
        request = urllib.request.Request(endpoint, data=json.dumps(request_payload).encode(), headers=headers, method="POST")
        with urllib.request.urlopen(request, timeout=120) as response: result = json.loads(response.read().decode("utf-8"))
        message = result.get("choices", [{}])[0].get("message", {})
        calls = message.get("tool_calls") or []
        if not calls: return message.get("content", ""), result.get("usage", {})
        messages.append(message)
        for call in calls:
            function = call.get("function", {}); name = function.get("name", "")
            try: arguments = json.loads(function.get("arguments") or "{}")
            except json.JSONDecodeError: arguments = {}
            output = execute_tool(tool_map.get(name, {"kind": "", "description": ""}), arguments)
            messages.append({"role": "tool", "tool_call_id": call.get("id", ""), "content": str(output)})
    raise ValueError("工具调用超过最大轮数")


def run_workflow_job(job_id, workflow_id, user_id, workflow_input):
    connection = db(); now = int(time.time())
    try:
        connection.execute("UPDATE ai_jobs SET status='running',updated_at=? WHERE id=?", (now, job_id)); connection.commit()
        workflow = connection.execute("SELECT * FROM ai_workflows WHERE id=? AND enabled=1", (workflow_id,)).fetchone()
        if not workflow: raise ValueError("工作流不存在或已停用")
        steps = json.loads(workflow["steps"] or "[]"); context = str(workflow_input); outputs = []
        for index, step in enumerate(steps):
            state = connection.execute("SELECT status FROM ai_jobs WHERE id=?", (job_id,)).fetchone()
            if state and state["status"] == "cancelled": return
            kind = str(step.get("type", "prompt"))
            if kind == "tool":
                tool = connection.execute("SELECT * FROM tools WHERE id=? AND enabled=1", (str(step.get("tool_id", "")),)).fetchone()
                if not tool: raise ValueError(f"第 {index + 1} 步工具不可用")
                arguments = step.get("arguments", {}) if isinstance(step.get("arguments", {}), dict) else {}
                if tool["kind"] == "calculator" and not arguments.get("expression"): arguments["expression"] = context
                result = execute_tool(dict(tool), arguments)
            elif kind == "search":
                query = str(step.get("query", "{{input}}")).replace("{{input}}", str(workflow_input)).replace("{{previous}}", context)
                result = json.dumps(web_search(query, max(1, min(int(step.get("limit", 5)), 10))), ensure_ascii=False)
            elif kind == "prompt":
                prompt_text = str(step.get("content", "{{input}}")).replace("{{input}}", str(workflow_input)).replace("{{previous}}", context)
                model = connection.execute("SELECT * FROM models WHERE id=? AND enabled=1", (str(step.get("model_id", "")),)).fetchone() if step.get("model_id") else connection.execute("SELECT * FROM models WHERE enabled=1 AND model_type='chat' ORDER BY is_default DESC,pinned DESC,name LIMIT 1").fetchone()
                if not model: raise ValueError("没有可用聊天模型")
                provider = connection.execute("SELECT * FROM provider_connections WHERE id=? AND enabled=1", (model["connection_id"],)).fetchone()
                endpoint = provider_api_base(provider["base_url"], provider["provider_type"]) + "/chat/completions"
                payload = {"model": model["base_model"], "messages": [{"role": "user", "content": prompt_text}], "stream": False}
                request = urllib.request.Request(endpoint, data=json.dumps(payload).encode(), headers={"Content-Type": "application/json", "Authorization": "Bearer " + provider["api_key"], "User-Agent": "RuoShopAdmin/1.0"}, method="POST")
                with urllib.request.urlopen(request, timeout=180) as response: response_data = json.loads(response.read().decode("utf-8"))
                result = str(response_data.get("choices", [{}])[0].get("message", {}).get("content", ""))
            else: raise ValueError(f"不支持的工作流步骤：{kind}")
            outputs.append({"step": index + 1, "type": kind, "output": result}); context = str(result)
        connection.execute("UPDATE ai_jobs SET status='completed',output=?,updated_at=? WHERE id=?", (json.dumps({"steps": outputs, "result": context}, ensure_ascii=False), int(time.time()), job_id)); connection.commit()
    except Exception as error:
        connection.execute("UPDATE ai_jobs SET status='failed',error=?,updated_at=? WHERE id=?", (str(error)[:2000], int(time.time()), job_id)); connection.commit()
    finally: connection.close()


def model_runtime_config(model):
    if not model:
        return [], []
    def decode(name):
        try:
            value = json.loads(model[name] or "[]")
            return [str(item).strip() for item in value if str(item).strip()] if isinstance(value, list) else []
        except (TypeError, json.JSONDecodeError):
            return []
    return decode("filters"), decode("actions")


def apply_input_filters(question, filters):
    value = question
    for name in filters:
        key = name.lower().replace("-", "_").replace(" ", "_")
        if key in ("trim", "trim_input"):
            value = value.strip()
        elif key in ("collapse_whitespace", "normalize_whitespace"):
            value = re.sub(r"[ \t]+", " ", value).strip()
        elif key in ("redact_secrets", "secret_guard"):
            value = re.sub(r"(?i)(api[_ -]?key|token|password)\s*[:=]\s*\S+", r"\1=[REDACTED]", value)
        elif key in ("max_12000", "limit_input"):
            value = value[:12000]
    return value


def action_instructions(actions):
    builtins = {
        "concise": "回答应简洁，优先给出可执行结论。",
        "cite_sources": "使用提供的资料时标注对应来源编号。",
        "structured": "复杂回答使用清晰的小标题和列表组织。",
        "json": "仅输出有效 JSON，不要添加 Markdown 代码围栏。",
        "translate_zh": "最终回答使用简体中文。",
    }
    instructions = [builtins.get(item.lower().replace("-", "_"), item) for item in actions]
    return "\n".join(instructions)


def index_file(connection, file_id, title, content, size=1000, overlap=150):
    old_ids = [row["id"] for row in connection.execute("SELECT id FROM file_chunks WHERE file_id=?", (file_id,))]
    for chunk_id in old_ids: connection.execute("DELETE FROM file_chunks_fts WHERE chunk_id=?", (chunk_id,))
    connection.execute("DELETE FROM file_chunks WHERE file_id=?", (file_id,))
    position, index = 0, 0
    while position < len(content):
        end = min(len(content), position + size); chunk = content[position:end].strip()
        if chunk:
            chunk_id = f"{file_id}-{index}"; connection.execute("INSERT INTO file_chunks(id,file_id,chunk_index,content) VALUES(?,?,?,?)", (chunk_id, file_id, index, chunk)); connection.execute("INSERT INTO file_chunks_fts(chunk_id,file_id,title,content) VALUES(?,?,?,?)", (chunk_id, file_id, title, chunk)); index += 1
        if end >= len(content): break
        position = end - overlap


def embedding_vectors(texts):
    model = setting("embedding_model")
    if not model or not setting("api_key") or not texts: return []
    endpoint = setting("base_url", "https://api.openai.com/v1").rstrip("/") + "/embeddings"
    request = urllib.request.Request(endpoint, data=json.dumps({"model": model, "input": texts}).encode(), headers={"Content-Type": "application/json", "Authorization": "Bearer " + setting("api_key")}, method="POST")
    with urllib.request.urlopen(request, timeout=90) as response: payload = json.loads(response.read().decode())
    return [item.get("embedding", []) for item in sorted(payload.get("data", []), key=lambda item: item.get("index", 0))]


def add_chunk_embeddings(connection, file_id):
    rows = connection.execute("SELECT id,content FROM file_chunks WHERE file_id=? ORDER BY chunk_index", (file_id,)).fetchall()
    try: vectors = embedding_vectors([row["content"] for row in rows])
    except Exception: return
    for row, vector in zip(rows, vectors):
        if vector: connection.execute("UPDATE file_chunks SET embedding=? WHERE id=?", (json.dumps(vector, separators=(",", ":")), row["id"]))


def fetch_provider_model_ids(provider=None):
    api_key = provider["api_key"] if provider else setting("api_key")
    provider_type = provider["provider_type"] if provider and "provider_type" in provider.keys() else "openai"
    manual_ids = parse_model_ids(provider["model_ids"] if provider and "model_ids" in provider.keys() else [])
    if manual_ids:
        return sorted(manual_ids)
    if not api_key and provider_type != "ollama": raise ValueError("请先在模型设置中配置 API Key")
    base_url = provider["base_url"] if provider else setting("base_url", "https://api.openai.com/v1")
    endpoint = base_url.rstrip("/") + ("/api/tags" if provider_type == "ollama" else "/models")
    headers = {"Accept": "application/json"}
    if api_key: headers["Authorization"] = "Bearer " + api_key
    request = urllib.request.Request(endpoint, headers=headers)
    try:
        with urllib.request.urlopen(request, timeout=30) as response:
            raw = response.read().decode("utf-8", errors="replace").lstrip("\ufeff \t\r\n")
    except urllib.error.HTTPError as error:
        raise upstream_http_error(error, "模型列表接口") from error
    try:
        payload = json.loads(raw)
    except json.JSONDecodeError as error:
        raise ValueError("接口地址返回的不是 JSON。请填写模型 API 接口地址，不要填写网页地址（例如 aistudio.google.com/app/apikey）。") from error
    items = payload.get("models", []) if provider_type == "ollama" else (payload.get("data", payload.get("models", [])) if isinstance(payload, dict) else [])
    return sorted({str(item.get("name", item.get("id", ""))).strip() for item in items if isinstance(item, dict) and str(item.get("name", item.get("id", ""))).strip()})


def sync_provider_models(connection, now, connection_id=""):
    provider = connection.execute("SELECT * FROM provider_connections WHERE id=?", (connection_id,)).fetchone() if connection_id else None
    model_ids = fetch_provider_model_ids(provider)
    for audio_model in provider_audio_presets(provider):
        if audio_model not in model_ids:
            model_ids.append(audio_model)
    if not model_ids:
        raise ValueError("模型接口没有返回可用模型")

    prefix = "model-provider-" + ((connection_id + "-") if connection_id else "")
    expected_ids = {prefix + hashlib.sha256(model_id.encode()).hexdigest()[:20] for model_id in model_ids}
    existing_ids = {row["id"] for row in connection.execute("SELECT id FROM models WHERE id LIKE ?", (prefix + "%",))}
    removed_ids = existing_ids - expected_ids
    if removed_ids:
        connection.executemany("DELETE FROM models WHERE id=?", [(item_id,) for item_id in removed_ids])

    added = 0
    for base_model in model_ids:
        item_id = prefix + hashlib.sha256(base_model.encode()).hexdigest()[:20]
        existed = connection.execute("SELECT 1 FROM models WHERE id=?", (item_id,)).fetchone()
        connection.execute(
            "INSERT INTO models(id,name,base_model,description,system_prompt,capabilities,updated_at,temperature,top_p,max_tokens,knowledge_id,skill_ids,tool_ids,connection_id,provider_id,enabled) "
            "VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?) ON CONFLICT(id) DO UPDATE SET name=excluded.name,base_model=excluded.base_model,connection_id=excluded.connection_id,provider_id=excluded.provider_id,updated_at=excluded.updated_at",
            (item_id, base_model, base_model, "从模型服务同步", "", "[]", now, 0.7, 1, 2048, "", "[]", "[]", connection_id, str(provider["provider_id"] if provider and "provider_id" in provider.keys() else "custom"), 1),
        )
        connection.execute("UPDATE models SET model_type=? WHERE id=?", (infer_model_type(base_model), item_id))
        if not existed:
            added += 1
    connection.commit()
    return {"total": len(model_ids), "added": added, "removed": len(removed_ids)}


def sync_enabled_provider_models(connection, now):
    providers = connection.execute("SELECT id FROM provider_connections WHERE enabled=1 ORDER BY updated_at DESC").fetchall()
    if not providers:
        raise ValueError("请先配置并启用模型连接")
    results = []
    errors = []
    for provider in providers:
        try:
            results.append(sync_provider_models(connection, now, provider["id"]))
        except Exception as error:
            errors.append({"id": provider["id"], "error": str(error)})
    return {
        "total": sum(item["total"] for item in results),
        "added": sum(item["added"] for item in results),
        "removed": sum(item["removed"] for item in results),
        "errors": errors,
    }


def cosine_similarity(left, right):
    if not left or len(left) != len(right): return 0.0
    dot = sum(a * b for a, b in zip(left, right)); norm = math.sqrt(sum(a * a for a in left) * sum(b * b for b in right))
    return dot / norm if norm else 0.0


def setting(key, default=""):
    connection = db()
    try:
        row = connection.execute("SELECT value FROM settings WHERE key=?", (key,)).fetchone()
        return row["value"] if row else default
    finally:
        connection.close()


def json_response(handler, status, payload):
    body = json.dumps(payload, ensure_ascii=False).encode("utf-8")
    handler.send_response(status)
    handler.send_header("Content-Type", "application/json; charset=utf-8")
    handler.send_header("Content-Length", str(len(body)))
    handler.end_headers()
    handler.wfile.write(body)


class Handler(BaseHTTPRequestHandler):
    def log_message(self, *_):
        return

    def read_json(self):
        size = int(self.headers.get("Content-Length", "0"))
        # A 15 MB file expands to roughly 20 MB after base64 encoding, with
        # additional JSON overhead. Keep the wire limit above the UI file cap.
        if size > 24_000_000:
            raise ValueError("请求内容过大")
        return json.loads(self.rfile.read(size).decode("utf-8")) if size else {}

    def identity(self):
        trusted_user = self.headers.get("X-Authenticated-User", "").strip()
        trusted_role = self.headers.get("X-Authenticated-Role", "").strip().lower()
        if trusted_user:
            return trusted_user, trusted_role if trusted_role in ("viewer", "editor", "admin", "superadmin") else "viewer"
        if REQUIRE_AUTH:
            raise PermissionError("AI 工作台身份验证失败")
        user_id = self.headers.get("X-Workspace-User", "local").strip() or "local"
        role = self.headers.get("X-Workspace-Role", "user").strip().lower()
        return user_id, role if role in ("user", "viewer", "editor", "admin", "superadmin") else "user"

    def require_role(self, minimum):
        levels = {"user": 0, "viewer": 0, "editor": 1, "admin": 2, "superadmin": 3}
        if levels.get(self.identity()[1], 0) < levels[minimum]:
            raise PermissionError("当前账号无权执行该操作")

    def enforce_rate_limit(self, connection, limit=60):
        user_id, role = self.identity()
        if role in ("admin", "superadmin"): return
        bucket = int(time.time() // 60)
        connection.execute("INSERT INTO ai_rate_limits(user_id,bucket,request_count) VALUES(?,?,1) ON CONFLICT(user_id,bucket) DO UPDATE SET request_count=request_count+1", (user_id, bucket))
        count = connection.execute("SELECT request_count FROM ai_rate_limits WHERE user_id=? AND bucket=?", (user_id, bucket)).fetchone()[0]
        connection.execute("DELETE FROM ai_rate_limits WHERE bucket<?", (bucket - 10,)); connection.commit()
        if count > limit: raise RateLimitError("请求过于频繁，请稍后再试")

    def can_access_model(self, row):
        user_id, role = self.identity()
        if role in ("admin", "superadmin") or not row:
            return True
        access = row["access"] if "access" in row.keys() else "private"
        if access == "public":
            return True
        try:
            grants = json.loads(row["access_grants"] or "[]")
        except (TypeError, json.JSONDecodeError):
            grants = []
        return (access == "private" and row["owner_id"] in ("", "local", user_id)) or (access == "shared" and (user_id in grants or role in grants))

    def do_GET(self):
        connection = None
        try:
            connection = db()
            if self.path == "/api/status":
                json_response(self, 200, {"online": True, "configured": bool(setting("api_key")), "model": setting("model", "gpt-4.1-mini"), "version": "ai-workspace-1.0"})
            elif self.path == "/api/config":
                json_response(self, 200, {"base_url": setting("base_url", "https://api.openai.com/v1"), "model": setting("model", "gpt-4.1-mini"), "embedding_model": setting("embedding_model"), "has_key": bool(setting("api_key"))})
            elif self.path == "/api/connections":
                rows = connection.execute("SELECT id,name,base_url,provider_type,provider_id,purpose,model_ids,enabled,updated_at,api_key FROM provider_connections ORDER BY updated_at DESC").fetchall()
                items = []
                for row in rows:
                    item = {**dict(row), "api_key": "", "model_ids": parse_model_ids(row["model_ids"]), "has_key": bool(row["api_key"]), "key_fingerprint": hashlib.sha256(row["api_key"].encode()).hexdigest()[:8].upper() if row["api_key"] else ""}
                    items.append(item)
                json_response(self, 200, {"connections": items})
            elif self.path == "/api/models":
                rows = connection.execute("SELECT * FROM models ORDER BY pinned DESC,sort_order,name").fetchall()
                json_response(self, 200, {"models": [dict(row) for row in rows if self.can_access_model(row)]})
            elif self.path == "/api/knowledge":
                json_response(self, 200, {"knowledge": [dict(row) for row in connection.execute("SELECT * FROM knowledge ORDER BY created_at DESC")]})
            elif self.path == "/api/files":
                rows = []
                for row in connection.execute("SELECT id,knowledge_id,name,status,created_at,source,metadata FROM files ORDER BY created_at DESC"):
                    item = dict(row)
                    try: item["image_count"] = len(json.loads(item.pop("metadata") or "{}").get("images", []))
                    except (ValueError, TypeError): item["image_count"] = 0
                    rows.append(item)
                json_response(self, 200, {"files": rows})
            elif self.path.startswith("/api/files/asset?"):
                query = urllib.parse.parse_qs(urllib.parse.urlparse(self.path).query)
                relative = urllib.parse.unquote(query.get("path", [""])[0]).replace("\\", "/").lstrip("/")
                target, root = (LEGACY_ASSETS / relative).resolve(), LEGACY_ASSETS.resolve()
                if not relative or not target.is_relative_to(root) or not target.is_file(): raise ValueError("图片不存在")
                raw = target.read_bytes(); self.send_response(200); self.send_header("Content-Type", mimetypes.guess_type(target.name)[0] or "application/octet-stream"); self.send_header("Content-Length", str(len(raw))); self.send_header("Cache-Control", "public, max-age=86400"); self.end_headers(); self.wfile.write(raw)
            elif self.path.startswith("/api/files/detail?"):
                file_id = urllib.parse.parse_qs(urllib.parse.urlparse(self.path).query).get("id", [""])[0]
                row = connection.execute("SELECT id,knowledge_id,name,content,status,created_at,source,metadata FROM files WHERE id=?", (file_id,)).fetchone()
                if not row: raise ValueError("文件不存在")
                item = dict(row)
                try: item["metadata"] = json.loads(item.get("metadata") or "{}")
                except (ValueError, TypeError): item["metadata"] = {}
                chunks = [dict(item) for item in connection.execute("SELECT id,chunk_index,content FROM file_chunks WHERE file_id=? ORDER BY chunk_index", (file_id,))]
                json_response(self, 200, {"file": item, "chunks": chunks})
            elif self.path == "/api/prompts":
                json_response(self, 200, {"prompts": [dict(row) for row in connection.execute("SELECT * FROM prompts ORDER BY updated_at DESC")]})
            elif self.path == "/api/notes":
                json_response(self, 200, {"notes": [dict(row) for row in connection.execute("SELECT * FROM notes WHERE owner_id=? ORDER BY updated_at DESC", (self.identity()[0],))]})
            elif self.path == "/api/skills":
                json_response(self, 200, {"skills": [dict(row) for row in connection.execute("SELECT * FROM skills ORDER BY updated_at DESC")]})
            elif self.path.startswith("/api/tools"):
                include_disabled = urllib.parse.parse_qs(urllib.parse.urlparse(self.path).query).get("all", [""])[0] == "1"
                where = "" if include_disabled else " WHERE enabled=1"
                json_response(self, 200, {"tools": [dict(row) for row in connection.execute("SELECT * FROM tools" + where + " ORDER BY name")]})
            elif self.path.startswith("/api/usage"):
                user_id, role = self.identity(); params = () if role in ("admin", "superadmin") else (user_id,); where = "" if not params else " WHERE user_id=?"
                rows = [dict(row) for row in connection.execute("SELECT * FROM ai_usage" + where + " ORDER BY created_at DESC LIMIT 500", params)]
                summary = connection.execute("SELECT COUNT(*) calls,COALESCE(SUM(input_tokens),0) input_tokens,COALESCE(SUM(output_tokens),0) output_tokens,COALESCE(SUM(cost),0) cost" + " FROM ai_usage" + where, params).fetchone()
                json_response(self, 200, {"usage": rows, "summary": dict(summary)})
            elif self.path == "/api/memories":
                user_id = self.identity()[0]; json_response(self, 200, {"memories": [dict(row) for row in connection.execute("SELECT * FROM ai_memories WHERE user_id=? ORDER BY updated_at DESC", (user_id,))]})
            elif self.path == "/api/jobs":
                user_id = self.identity()[0]; json_response(self, 200, {"jobs": [dict(row) for row in connection.execute("SELECT * FROM ai_jobs WHERE user_id=? ORDER BY updated_at DESC LIMIT 100", (user_id,))]})
            elif self.path == "/api/workflows":
                user_id, role = self.identity(); rows = connection.execute("SELECT * FROM ai_workflows WHERE owner_id=? OR ? IN ('admin','superadmin') ORDER BY updated_at DESC", (user_id, role)); json_response(self, 200, {"workflows": [dict(row) for row in rows]})
            elif self.path.startswith("/api/chats?") or self.path == "/api/chats":
                user_id = self.identity()[0]
                rows = []
                for row in connection.execute("SELECT * FROM chats WHERE user_id=? ORDER BY updated_at DESC LIMIT 100", (user_id,)):
                    item = dict(row); item["messages"] = json.loads(item["messages"]); rows.append(item)
                json_response(self, 200, {"chats": rows})
            elif self.path.startswith("/api/shares?"):
                share_id = urllib.parse.parse_qs(urllib.parse.urlparse(self.path).query).get("id", [""])[0]
                row = connection.execute("SELECT id,title,messages,created_at,expires_at,revoked FROM shared_chats WHERE id=?", (share_id,)).fetchone()
                if not row or row["revoked"] or (row["expires_at"] and row["expires_at"] < int(time.time())): raise ValueError("分享不存在或已失效")
                item = dict(row); item["messages"] = json.loads(item["messages"]); json_response(self, 200, {"share": item})
            elif self.path == "/api/shares":
                json_response(self, 200, {"shares": [dict(row) for row in connection.execute("SELECT id,title,created_at,expires_at,revoked FROM shared_chats WHERE owner_id=? ORDER BY created_at DESC", (self.identity()[0],))]})
            elif self.path == "/api/search":
                json_response(self, 200, {"documents": []})
            else:
                json_response(self, 404, {"error": "接口不存在"})
        except PermissionError as error:
            json_response(self, 401, {"error": str(error)})
        except RateLimitError as error:
            json_response(self, 429, {"error": str(error), "retry_after": 60})
        except Exception as error:
            json_response(self, 400, {"error": str(error)})
        finally:
            if connection is not None:
                connection.close()

    def do_POST(self):
        connection = None
        try:
            data = self.read_json(); connection = db(); now = int(time.time()); request_started = time.perf_counter()
            admin_paths = ("/api/connections/save", "/api/connections/delete", "/api/connections/sync", "/api/connections/models", "/api/connections/test", "/api/connections/toggle", "/api/config", "/api/models", "/api/models/sync", "/api/models/update", "/api/models/delete", "/api/prompts", "/api/prompts/delete", "/api/prompts/update", "/api/skills", "/api/skills/delete", "/api/skills/update", "/api/tools", "/api/tools/delete", "/api/tools/update", "/api/tools/test")
            editor_paths = ("/api/knowledge", "/api/knowledge/delete", "/api/documents/import-file", "/api/files/assign", "/api/files/reprocess", "/api/files/delete")
            if self.path in admin_paths: self.require_role("admin")
            elif self.path in editor_paths: self.require_role("editor")
            if self.path in ("/api/chat", "/api/chat/stream", "/api/images/generations", "/api/audio/transcriptions", "/api/audio/speech", "/api/web-search", "/api/files/generate", "/api/files/reprocess", "/api/tools/test"):
                self.enforce_rate_limit(connection)
            if self.path == "/api/connections/save":
                item_id = str(data.get("id", "")).strip() or "conn-" + hashlib.sha256(f"{data.get('base_url')}:{now}".encode()).hexdigest()[:16]
                name, base_url = str(data.get("name", "")).strip(), str(data.get("base_url", "")).strip().rstrip("/")
                if not name or not base_url: raise ValueError("连接名称和接口地址不能为空")
                old = connection.execute("SELECT api_key FROM provider_connections WHERE id=?", (item_id,)).fetchone()
                api_key = str(data.get("api_key", "")).strip() or (old["api_key"] if old else "")
                provider_type = str(data.get("provider_type", "openai")).strip().lower()
                if provider_type not in ("openai", "ollama", "pipeline"): provider_type = "openai"
                provider_id = str(data.get("provider_id", "custom")).strip().lower() or "custom"
                model_ids = parse_model_ids(data.get("model_ids", []))
                purpose = str(data.get("purpose", "general")).strip().lower()
                if purpose not in ("general", "chat", "image", "audio"): purpose = "general"
                if not api_key and provider_type != "ollama": raise ValueError("API Key 不能为空")
                connection.execute("INSERT INTO provider_connections(id,name,base_url,api_key,enabled,provider_type,provider_id,purpose,model_ids,updated_at) VALUES(?,?,?,?,?,?,?,?,?,?) ON CONFLICT(id) DO UPDATE SET name=excluded.name,base_url=excluded.base_url,api_key=excluded.api_key,enabled=excluded.enabled,provider_type=excluded.provider_type,provider_id=excluded.provider_id,purpose=excluded.purpose,model_ids=excluded.model_ids,updated_at=excluded.updated_at", (item_id, name, base_url, api_key, 1 if data.get("enabled", True) else 0, provider_type, provider_id, purpose, json.dumps(model_ids, ensure_ascii=False), now))
                connection.commit()
                result = {"total": 0, "added": 0, "removed": 0}
                sync_error = ""
                if data.get("sync_models", True):
                    try:
                        result = sync_provider_models(connection, now, item_id)
                    except Exception as error:
                        sync_error = str(error)
                json_response(self, 200, {"ok": True, "id": item_id, "sync": result, "sync_error": sync_error})
            elif self.path == "/api/connections/delete":
                item_id = str(data.get("id", "")); connection.execute("DELETE FROM models WHERE connection_id=?", (item_id,)); connection.execute("DELETE FROM provider_connections WHERE id=?", (item_id,)); connection.commit(); json_response(self, 200, {"ok": True})
            elif self.path == "/api/connections/sync":
                item_id = str(data.get("id", "")); result = sync_provider_models(connection, now, item_id); json_response(self, 200, {"ok": True, **result})
            elif self.path == "/api/connections/models":
                item_id = str(data.get("id", "")); provider = connection.execute("SELECT * FROM provider_connections WHERE id=? AND enabled=1", (item_id,)).fetchone()
                if not provider: raise ValueError("连接不存在或未启用")
                json_response(self, 200, {"ok": True, "models": fetch_provider_model_ids(provider)})
            elif self.path == "/api/connections/test":
                item_id = str(data.get("id", "")); provider = connection.execute("SELECT * FROM provider_connections WHERE id=?", (item_id,)).fetchone()
                if not provider: raise ValueError("连接不存在")
                model_ids = fetch_provider_model_ids(provider)
                json_response(self, 200, {"ok": True, "models": len(model_ids), "message": f"连接成功，发现 {len(model_ids)} 个模型"})
            elif self.path == "/api/connections/toggle":
                item_id = str(data.get("id", "")); enabled = 1 if data.get("enabled", True) else 0
                connection.execute("UPDATE provider_connections SET enabled=?,updated_at=? WHERE id=?", (enabled, now, item_id)); connection.execute("UPDATE models SET enabled=? WHERE connection_id=?", (enabled, item_id)); connection.commit(); json_response(self, 200, {"ok": True, "enabled": bool(enabled)})
            elif self.path == "/api/config":
                current_key_row = connection.execute("SELECT value FROM settings WHERE key='api_key'").fetchone()
                current_key = current_key_row["value"] if current_key_row else ""
                for key in ("base_url", "model", "embedding_model", "api_key"):
                    if key in data:
                        value = str(data[key]).strip()
                        # The UI never reads the stored secret back. An empty
                        # field therefore means "keep it", rather than erase it.
                        if key == "api_key" and not value and current_key:
                            continue
                        connection.execute("INSERT INTO settings(key,value) VALUES(?,?) ON CONFLICT(key) DO UPDATE SET value=excluded.value", (key, value))
                connection.commit()
                sync_result = sync_enabled_provider_models(connection, now) if data.get("sync_models") and connection.execute("SELECT 1 FROM provider_connections WHERE enabled=1").fetchone() else {"total": 0, "added": 0, "removed": 0}
                json_response(self, 200, {"ok": True, "base_url": setting("base_url"), "model": setting("model"), "embedding_model": setting("embedding_model"), "has_key": bool(setting("api_key")), "sync": sync_result})
            elif self.path == "/api/memories":
                content = str(data.get("content", "")).strip()
                if not content: raise ValueError("记忆内容不能为空")
                item_id = str(data.get("id", ""))
                if item_id:
                    connection.execute("UPDATE ai_memories SET content=?,source_chat_id=?,enabled=?,updated_at=? WHERE id=? AND user_id=?", (content[:8000], str(data.get("source_chat_id", "")), 1 if data.get("enabled", True) else 0, now, item_id, self.identity()[0]))
                else:
                    item_id = "memory-" + hashlib.sha256(f"{self.identity()[0]}:{content}:{now}".encode()).hexdigest()[:20]
                    connection.execute("INSERT INTO ai_memories(id,user_id,content,source_chat_id,enabled,created_at,updated_at) VALUES(?,?,?,?,?,?,?)", (item_id, self.identity()[0], content[:8000], str(data.get("source_chat_id", "")), 1 if data.get("enabled", True) else 0, now, now))
                connection.commit(); json_response(self, 200, {"ok": True, "id": item_id})
            elif self.path == "/api/memories/delete":
                connection.execute("DELETE FROM ai_memories WHERE id=? AND user_id=?", (str(data.get("id", "")), self.identity()[0])); connection.commit(); json_response(self, 200, {"ok": True})
            elif self.path == "/api/workflows":
                name = str(data.get("name", "")).strip(); steps = data.get("steps", [])
                if not name or not isinstance(steps, list) or not steps: raise ValueError("工作流至少需要一个步骤")
                item_id = str(data.get("id", ""))
                if item_id:
                    connection.execute("UPDATE ai_workflows SET name=?,description=?,steps=?,enabled=?,updated_at=? WHERE id=? AND owner_id=?", (name, str(data.get("description", "")), json.dumps(steps, ensure_ascii=False), 1 if data.get("enabled", True) else 0, now, item_id, self.identity()[0]))
                else:
                    item_id = "workflow-" + hashlib.sha256(f"{name}:{now}".encode()).hexdigest()[:20]
                    connection.execute("INSERT INTO ai_workflows(id,owner_id,name,description,steps,enabled,created_at,updated_at) VALUES(?,?,?,?,?,?,?,?)", (item_id, self.identity()[0], name, str(data.get("description", "")), json.dumps(steps, ensure_ascii=False), 1 if data.get("enabled", True) else 0, now, now))
                connection.commit(); json_response(self, 200, {"ok": True, "id": item_id})
            elif self.path == "/api/workflows/delete":
                connection.execute("DELETE FROM ai_workflows WHERE id=? AND owner_id=?", (str(data.get("id", "")), self.identity()[0])); connection.commit(); json_response(self, 200, {"ok": True})
            elif self.path == "/api/workflows/run":
                workflow_id = str(data.get("id", "")); workflow = connection.execute("SELECT id FROM ai_workflows WHERE id=? AND enabled=1 AND owner_id=?", (workflow_id, self.identity()[0])).fetchone()
                if not workflow: raise ValueError("工作流不存在或已停用")
                job_id = "job-" + hashlib.sha256(f"{workflow_id}:{self.identity()[0]}:{time.time_ns()}".encode()).hexdigest()[:24]
                connection.execute("INSERT INTO ai_jobs(id,user_id,kind,status,input,created_at,updated_at) VALUES(?,?,?,?,?,?,?)", (job_id, self.identity()[0], "workflow", "queued", json.dumps({"workflow_id": workflow_id, "input": data.get("input", "")}, ensure_ascii=False), now, now)); connection.commit()
                threading.Thread(target=run_workflow_job, args=(job_id, workflow_id, self.identity()[0], data.get("input", "")), daemon=True).start(); json_response(self, 202, {"ok": True, "job_id": job_id, "status": "queued"})
            elif self.path == "/api/jobs/delete":
                connection.execute("DELETE FROM ai_jobs WHERE id=? AND user_id=? AND status IN ('completed','failed','cancelled')", (str(data.get("id", "")), self.identity()[0])); connection.commit(); json_response(self, 200, {"ok": True})
            elif self.path == "/api/jobs/cancel":
                connection.execute("UPDATE ai_jobs SET status='cancelled',updated_at=? WHERE id=? AND user_id=? AND status IN ('queued','running')", (now, str(data.get("id", "")), self.identity()[0])); connection.commit(); json_response(self, 200, {"ok": True})
            elif self.path == "/api/jobs/retry":
                row = connection.execute("SELECT input FROM ai_jobs WHERE id=? AND user_id=? AND status IN ('completed','failed','cancelled')", (str(data.get("id", "")), self.identity()[0])).fetchone()
                if not row: raise ValueError("任务不可重试")
                job_input = json.loads(row["input"] or "{}"); workflow_id = str(job_input.get("workflow_id", "")); workflow_input = job_input.get("input", "")
                if not connection.execute("SELECT 1 FROM ai_workflows WHERE id=? AND owner_id=? AND enabled=1", (workflow_id, self.identity()[0])).fetchone(): raise ValueError("工作流不存在或已停用")
                job_id = "job-" + hashlib.sha256(f"{workflow_id}:{self.identity()[0]}:{time.time_ns()}".encode()).hexdigest()[:24]
                connection.execute("INSERT INTO ai_jobs(id,user_id,kind,status,input,created_at,updated_at) VALUES(?,?,?,?,?,?,?)", (job_id, self.identity()[0], "workflow", "queued", json.dumps(job_input, ensure_ascii=False), now, now)); connection.commit(); threading.Thread(target=run_workflow_job, args=(job_id, workflow_id, self.identity()[0], workflow_input), daemon=True).start(); json_response(self, 202, {"ok": True, "job_id": job_id})
            elif self.path == "/api/knowledge":
                name = str(data.get("name", "")).strip()
                if not name: raise ValueError("知识集合名称不能为空")
                item_id = "kb-" + hashlib.sha256(f"{name}:{now}".encode()).hexdigest()[:20]
                connection.execute("INSERT INTO knowledge(id,name,description,created_at) VALUES(?,?,?,?)", (item_id, name, str(data.get("description", "")), now)); connection.commit(); json_response(self, 201, {"ok": True, "id": item_id})
            elif self.path == "/api/models":
                name = str(data.get("name", "")).strip(); base_model = str(data.get("base_model", "")).strip()
                if not name or not base_model: raise ValueError("模型名称和基础模型不能为空")
                connection_id = str(data.get("connection_id", ""))
                if connection_id and not connection.execute("SELECT 1 FROM provider_connections WHERE id=?", (connection_id,)).fetchone():
                    raise ValueError("所属账号连接不存在")
                item_id = "model-" + hashlib.sha256(f"{name}:{now}".encode()).hexdigest()[:20]
                capabilities = json.dumps(data.get("capabilities", ["knowledge"]), ensure_ascii=False)
                temperature = max(0, min(float(data.get("temperature", 0.7)), 2)); top_p = max(0, min(float(data.get("top_p", 1)), 1)); max_tokens = max(1, min(int(data.get("max_tokens", 2048)), 128000))
                enabled = 1 if data.get("enabled", True) else 0; hidden = 1 if data.get("hidden", False) else 0; pinned = 1 if data.get("pinned", False) else 0; is_default = 1 if data.get("is_default", False) else 0
                if is_default: connection.execute("UPDATE models SET is_default=0")
                owner_id = self.identity()[0]
                connection.execute("INSERT INTO models(id,name,base_model,description,system_prompt,capabilities,updated_at,temperature,top_p,max_tokens,knowledge_id,skill_ids,tool_ids,connection_id,enabled,hidden,pinned,is_default,tags,sort_order,owner_id) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)", (item_id, name, base_model, str(data.get("description", "")), str(data.get("system_prompt", "")), capabilities, now, temperature, top_p, max_tokens, str(data.get("knowledge_id", "")), json.dumps(data.get("skill_ids", [])), json.dumps(data.get("tool_ids", [])), connection_id, enabled, hidden, pinned, is_default, json.dumps(data.get("tags", []), ensure_ascii=False), int(data.get("sort_order", 0)), owner_id))
                connection.execute("UPDATE models SET provider_id=COALESCE((SELECT provider_id FROM provider_connections WHERE id=?),provider_id) WHERE id=?", (connection_id, item_id))
                connection.execute("UPDATE models SET model_type=?,input_price=?,output_price=? WHERE id=?", (str(data.get("model_type") or infer_model_type(base_model)), max(0, float(data.get("input_price", 0))), max(0, float(data.get("output_price", 0))), item_id)); connection.commit(); json_response(self, 201, {"ok": True, "id": item_id})
            elif self.path == "/api/models/sync":
                result = sync_enabled_provider_models(connection, now)
                response = {"ok": True, **result}
                if not response.get("removed"): response.pop("removed", None)
                json_response(self, 200, response)
            elif self.path == "/api/models/update":
                item_id, name, base_model = str(data.get("id", "")), str(data.get("name", "")).strip(), str(data.get("base_model", "")).strip()
                if not connection.execute("SELECT 1 FROM models WHERE id=?", (item_id,)).fetchone(): raise ValueError("模型不存在")
                if not name or not base_model: raise ValueError("模型名称和基础模型不能为空")
                temperature = max(0, min(float(data.get("temperature", 0.7)), 2)); top_p = max(0, min(float(data.get("top_p", 1)), 1)); max_tokens = max(1, min(int(data.get("max_tokens", 2048)), 128000))
                enabled = 1 if data.get("enabled", True) else 0; hidden = 1 if data.get("hidden", False) else 0; pinned = 1 if data.get("pinned", False) else 0; is_default = 1 if data.get("is_default", False) else 0
                if is_default: connection.execute("UPDATE models SET is_default=0")
                capabilities = json.dumps(data.get("capabilities", []), ensure_ascii=False)
                connection_id = str(data.get("connection_id", ""))
                if connection_id and not connection.execute("SELECT 1 FROM provider_connections WHERE id=?", (connection_id,)).fetchone():
                    raise ValueError("所属账号连接不存在")
                connection.execute("UPDATE models SET name=?,base_model=?,description=?,system_prompt=?,capabilities=?,temperature=?,top_p=?,max_tokens=?,knowledge_id=?,skill_ids=?,tool_ids=?,connection_id=?,provider_id=COALESCE((SELECT provider_id FROM provider_connections WHERE id=?),provider_id),enabled=?,hidden=?,pinned=?,is_default=?,tags=?,sort_order=?,access=?,access_grants=?,filters=?,actions=?,updated_at=? WHERE id=?", (name, base_model, str(data.get("description", "")), str(data.get("system_prompt", "")), capabilities, temperature, top_p, max_tokens, str(data.get("knowledge_id", "")), json.dumps(data.get("skill_ids", [])), json.dumps(data.get("tool_ids", [])), connection_id, connection_id, enabled, hidden, pinned, is_default, json.dumps(data.get("tags", []), ensure_ascii=False), int(data.get("sort_order", 0)), str(data.get("access", "private")), json.dumps(data.get("access_grants", []), ensure_ascii=False), json.dumps(data.get("filters", []), ensure_ascii=False), json.dumps(data.get("actions", []), ensure_ascii=False), now, item_id))
                connection.execute("UPDATE models SET model_type=?,input_price=?,output_price=? WHERE id=?", (str(data.get("model_type") or infer_model_type(base_model)), max(0, float(data.get("input_price", 0))), max(0, float(data.get("output_price", 0))), item_id)); connection.commit(); json_response(self, 200, {"ok": True, "id": item_id, "enabled": bool(enabled)})
            elif self.path == "/api/tools":
                name, description, kind = str(data.get("name", "")).strip(), str(data.get("description", "")).strip(), str(data.get("kind", "custom")).strip()
                if not name: raise ValueError("工具名称不能为空")
                item_id = "tool-" + hashlib.sha256(f"{name}:{now}".encode()).hexdigest()[:20]
                connection.execute("INSERT INTO tools(id,name,description,kind,enabled,updated_at,config) VALUES(?,?,?,?,?,?,?)", (item_id, name, description, kind, 1 if data.get("enabled", True) else 0, now, json.dumps(data.get("config", {}), ensure_ascii=False))); connection.commit(); json_response(self, 201, {"ok": True, "id": item_id})
            elif self.path == "/api/prompts":
                command, title, content = str(data.get("command", "")).strip().lstrip("/"), str(data.get("title", "")).strip(), str(data.get("content", ""))
                if not command or not title or not content: raise ValueError("Prompt 信息不完整")
                item_id = "prompt-" + hashlib.sha256(f"{command}:{now}".encode()).hexdigest()[:20]
                connection.execute("INSERT OR REPLACE INTO prompts(id,command,title,content,updated_at) VALUES(?,?,?,?,?)", (item_id, command, title, content, now)); connection.commit(); json_response(self, 201, {"ok": True, "id": item_id})
            elif self.path == "/api/skills":
                name, content = str(data.get("name", "")).strip(), str(data.get("content", "")).strip()
                if not name or not content: raise ValueError("技能名称和内容不能为空")
                item_id = "skill-" + hashlib.sha256(f"{name}:{now}".encode()).hexdigest()[:20]
                connection.execute("INSERT INTO skills(id,name,description,content,updated_at) VALUES(?,?,?,?,?)", (item_id, name, str(data.get("description", "")), content, now)); connection.commit(); json_response(self, 201, {"ok": True, "id": item_id})
            elif self.path == "/api/chats/save":
                chat_id, user_id, title = str(data.get("id", "")).strip(), self.identity()[0], str(data.get("title", "新对话"))[:200]
                messages = data.get("messages", [])
                if not chat_id or not isinstance(messages, list): raise ValueError("会话数据不正确")
                created_at = int(data.get("created_at") or now)
                connection.execute("INSERT INTO chats(id,user_id,title,messages,folder,archived,model_id,favorite,parent_chat_id,created_at,updated_at) VALUES(?,?,?,?,?,?,?,?,?,?,?) ON CONFLICT(id) DO UPDATE SET title=excluded.title,messages=excluded.messages,folder=excluded.folder,archived=excluded.archived,model_id=excluded.model_id,favorite=excluded.favorite,parent_chat_id=excluded.parent_chat_id,updated_at=excluded.updated_at", (chat_id, user_id, title, json.dumps(messages, ensure_ascii=False), str(data.get("folder", "")), 1 if data.get("archived", False) else 0, str(data.get("model_id", "")), 1 if data.get("favorite", False) else 0, str(data.get("parent_chat_id", "")), created_at, now)); connection.commit(); json_response(self, 200, {"ok": True})
            elif self.path == "/api/chats/delete":
                connection.execute("DELETE FROM chats WHERE id=? AND user_id=?", (str(data.get("id", "")), self.identity()[0])); connection.commit(); json_response(self, 200, {"ok": True})
            elif self.path == "/api/chats/search":
                query = str(data.get("query", "")).strip()
                if not query: json_response(self, 200, {"results": []}); return
                pattern = f"%{query}%"; matches = []
                for row in connection.execute("SELECT id,title,messages,updated_at FROM chats WHERE user_id=? AND (title LIKE ? OR messages LIKE ?) ORDER BY updated_at DESC LIMIT 50", (self.identity()[0], pattern, pattern)):
                    messages = json.loads(row["messages"] or "[]"); snippet = ""
                    for message in messages:
                        content = str(message.get("content", "")); position = content.lower().find(query.lower())
                        if position >= 0: snippet = content[max(0, position - 60):position + len(query) + 100]; break
                    matches.append({"id": row["id"], "title": row["title"], "snippet": snippet, "updated_at": row["updated_at"]})
                json_response(self, 200, {"results": matches})
            elif self.path == "/api/shares":
                title, messages = str(data.get("title", "共享会话"))[:200], data.get("messages", [])
                if not isinstance(messages, list): raise ValueError("会话数据不正确")
                share_id = hashlib.sha256(f"{self.identity()[0]}:{now}:{json.dumps(messages, ensure_ascii=False)}".encode()).hexdigest()[:24]
                expires_at = now + max(1, min(int(data.get("expires_in_days", 7)), 365)) * 86400
                connection.execute("INSERT INTO shared_chats(id,owner_id,title,messages,created_at,expires_at,revoked) VALUES(?,?,?,?,?,?,0)", (share_id, self.identity()[0], title, json.dumps(messages, ensure_ascii=False), now, expires_at)); connection.commit(); json_response(self, 201, {"ok": True, "id": share_id, "expires_at": expires_at})
            elif self.path == "/api/shares/revoke":
                connection.execute("UPDATE shared_chats SET revoked=1 WHERE id=? AND owner_id=?", (str(data.get("id", "")), self.identity()[0])); connection.commit(); json_response(self, 200, {"ok": True})
            elif self.path == "/api/documents/import-file":
                name = Path(str(data.get("filename", "")).strip()).name
                raw = base64.b64decode(str(data.get("data", "")), validate=True)
                if not name or not raw: raise ValueError("文件不能为空")
                if len(raw) > 15_000_000: raise ValueError("单个文件不能超过 15MB")
                content, status = extract_text(name, raw)
                FILES.mkdir(parents=True, exist_ok=True); file_id = "file-" + hashlib.sha256(raw).hexdigest()[:24]; path = FILES / f"{file_id}-{name}"; path.write_bytes(raw)
                connection.execute("INSERT OR REPLACE INTO files(id,name,content,path,status,created_at) VALUES(?,?,?,?,?,?)", (file_id, name, content, str(path), status, now))
                if content: index_file(connection, file_id, name, content); add_chunk_embeddings(connection, file_id)
                connection.commit(); json_response(self, 201, {"ok": True, "file": {"id": file_id, "name": name, "status": status}})
            elif self.path == "/api/files/assign":
                file_id, knowledge_id = str(data.get("file_id", "")), str(data.get("knowledge_id", "")) or None
                if not connection.execute("SELECT 1 FROM files WHERE id=?", (file_id,)).fetchone(): raise ValueError("文件不存在")
                if knowledge_id and not connection.execute("SELECT 1 FROM knowledge WHERE id=?", (knowledge_id,)).fetchone(): raise ValueError("知识集合不存在")
                connection.execute("UPDATE files SET knowledge_id=? WHERE id=?", (knowledge_id, file_id)); connection.commit(); json_response(self, 200, {"ok": True})
            elif self.path == "/api/files/reprocess":
                file_id = str(data.get("id", "")); row = connection.execute("SELECT * FROM files WHERE id=?", (file_id,)).fetchone()
                if not row: raise ValueError("文件不存在")
                path = Path(row["path"]); raw = path.read_bytes()
                if path.suffix.lower() in (".png", ".jpg", ".jpeg", ".webp"):
                    content = ocr_image(connection, row["name"], raw, str(data.get("model_id", "")))
                else:
                    content, _ = extract_text(row["name"], raw)
                for chunk in connection.execute("SELECT id FROM file_chunks WHERE file_id=?", (file_id,)): connection.execute("DELETE FROM file_chunks_fts WHERE chunk_id=?", (chunk["id"],))
                connection.execute("DELETE FROM file_chunks WHERE file_id=?", (file_id,)); connection.execute("UPDATE files SET content=?,status='ready' WHERE id=?", (content, file_id))
                index_file(connection, file_id, row["name"], content); add_chunk_embeddings(connection, file_id); connection.commit()
                json_response(self, 200, {"ok": True, "status": "ready", "characters": len(content)})
            elif self.path == "/api/files/delete":
                row = connection.execute("SELECT path FROM files WHERE id=?", (str(data.get("id", "")),)).fetchone()
                if row and row["path"]:
                    path = Path(row["path"]).resolve(); root = FILES.resolve()
                    if path.is_relative_to(root) and path.exists(): path.unlink()
                for chunk in connection.execute("SELECT id FROM file_chunks WHERE file_id=?", (str(data.get("id", "")),)): connection.execute("DELETE FROM file_chunks_fts WHERE chunk_id=?", (chunk["id"],))
                connection.execute("DELETE FROM file_chunks WHERE file_id=?", (str(data.get("id", "")),))
                connection.execute("DELETE FROM files WHERE id=?", (str(data.get("id", "")),)); connection.commit(); json_response(self, 200, {"ok": True})
            elif self.path == "/api/knowledge/delete":
                item_id = str(data.get("id", "")); connection.execute("UPDATE files SET knowledge_id=NULL WHERE knowledge_id=?", (item_id,)); connection.execute("DELETE FROM knowledge WHERE id=?", (item_id,)); connection.commit(); json_response(self, 200, {"ok": True})
            elif self.path == "/api/models/delete":
                connection.execute("DELETE FROM models WHERE id=?", (str(data.get("id", "")),)); connection.commit(); json_response(self, 200, {"ok": True})
            elif self.path == "/api/prompts/delete":
                connection.execute("DELETE FROM prompts WHERE id=?", (str(data.get("id", "")),)); connection.commit(); json_response(self, 200, {"ok": True})
            elif self.path == "/api/notes/delete":
                connection.execute("DELETE FROM notes WHERE id=? AND owner_id=?", (str(data.get("id", "")), self.identity()[0])); connection.commit(); json_response(self, 200, {"ok": True})
            elif self.path == "/api/prompts/update":
                item_id = str(data.get("id", "")); command = str(data.get("command", "")).strip().lstrip("/"); title = str(data.get("title", "")).strip(); content = str(data.get("content", ""))
                if not item_id or not command or not title or not content: raise ValueError("Prompt 信息不完整")
                connection.execute("UPDATE prompts SET command=?,title=?,content=?,updated_at=? WHERE id=?", (command, title, content, now, item_id)); connection.commit(); json_response(self, 200, {"ok": True})
            elif self.path == "/api/notes":
                title, content = str(data.get("title", "")).strip(), str(data.get("content", ""))
                if not title or not content: raise ValueError("笔记信息不完整")
                item_id = "note-" + hashlib.sha256(f"{title}:{now}".encode()).hexdigest()[:20]
                connection.execute("INSERT INTO notes(id,title,content,updated_at,owner_id) VALUES(?,?,?,?,?)", (item_id, title, content, now, self.identity()[0])); connection.commit(); json_response(self, 201, {"ok": True, "id": item_id})
            elif self.path == "/api/notes/update":
                item_id = str(data.get("id", "")); title, content = str(data.get("title", "")).strip(), str(data.get("content", ""))
                if not item_id or not title or not content: raise ValueError("笔记信息不完整")
                connection.execute("UPDATE notes SET title=?,content=?,updated_at=? WHERE id=? AND owner_id=?", (title, content, now, item_id, self.identity()[0])); connection.commit(); json_response(self, 200, {"ok": True})
            elif self.path == "/api/skills/update":
                item_id = str(data.get("id", "")); name = str(data.get("name", "")).strip(); content = str(data.get("content", "")).strip()
                if not item_id or not name or not content: raise ValueError("Skill 信息不完整")
                connection.execute("UPDATE skills SET name=?,description=?,content=?,updated_at=? WHERE id=?", (name, str(data.get("description", "")), content, now, item_id)); connection.commit(); json_response(self, 200, {"ok": True})
            elif self.path == "/api/tools/update":
                item_id = str(data.get("id", "")); name = str(data.get("name", "")).strip()
                if not item_id or not name: raise ValueError("工具信息不完整")
                connection.execute("UPDATE tools SET name=?,description=?,kind=?,enabled=?,updated_at=?,config=? WHERE id=? AND id NOT LIKE 'builtin-%'", (name, str(data.get("description", "")), str(data.get("kind", "custom")), 1 if data.get("enabled", True) else 0, now, json.dumps(data.get("config", {}), ensure_ascii=False), item_id)); connection.commit(); json_response(self, 200, {"ok": True})
            elif self.path == "/api/tools/test":
                row = {"kind": str(data.get("kind", "http")), "description": str(data.get("description", "")), "config": json.dumps(data.get("config", {}), ensure_ascii=False)}
                json_response(self, 200, {"ok": True, "result": execute_tool(row, data.get("arguments", {}))[:200000]})
            elif self.path == "/api/tools/delete":
                connection.execute("DELETE FROM tools WHERE id=? AND id NOT LIKE 'builtin-%'", (str(data.get("id", "")),)); connection.commit(); json_response(self, 200, {"ok": True})
            elif self.path == "/api/skills/delete":
                connection.execute("DELETE FROM skills WHERE id=?", (str(data.get("id", "")),)); connection.commit(); json_response(self, 200, {"ok": True})
            elif self.path == "/api/search":
                query = str(data.get("query", "")).lower()
                knowledge_id = str(data.get("knowledge_id", "")).strip()
                params = [query]; where = "file_chunks_fts MATCH ?"
                if knowledge_id: where += " AND f.knowledge_id=?"; params.append(knowledge_id)
                try:
                    rows = connection.execute(f"SELECT c.id AS chunk_id,f.id,f.name AS title,c.content,f.status,bm25(file_chunks_fts) AS rank FROM file_chunks_fts JOIN file_chunks c ON c.id=file_chunks_fts.chunk_id JOIN files f ON f.id=c.file_id WHERE {where} ORDER BY rank LIMIT 10", params).fetchall()
                except sqlite3.OperationalError:
                    rows = []
                # Trigram FTS does not match very short terms. The LIKE fallback
                # also covers punctuation-heavy input that FTS cannot parse.
                if not rows:
                    cleaned = re.sub(r"帮我|请问|看一下|查一下|介绍一下|告诉我|有几种|有哪些|是什么|怎么|如何|为什么|一下", " ", query)
                    terms = re.findall(r"[a-z0-9][a-z0-9._-]+|[\u4e00-\u9fff]{2,}", cleaned)
                    expanded = []
                    for term in terms:
                        if re.fullmatch(r"[\u4e00-\u9fff]{5,}", term): expanded.extend(term[index:index + 2] for index in range(0, len(term) - 1, 2))
                        else: expanded.append(term)
                    terms = list(dict.fromkeys(term for term in expanded if len(term) >= 2))[:8] or [query]
                    clauses = ["(lower(f.name) LIKE ? OR lower(c.content) LIKE ?)" for _ in terms]
                    params = [value for term in terms for value in (f"%{term}%", f"%{term}%")]
                    scope = " AND f.knowledge_id=?" if knowledge_id else ""
                    if knowledge_id: params.append(knowledge_id)
                    rows = connection.execute(f"SELECT c.id AS chunk_id,f.id,f.name AS title,c.content,f.status,0 AS rank FROM file_chunks c JOIN files f ON f.id=c.file_id WHERE ({' OR '.join(clauses)}){scope} LIMIT 50", params).fetchall()
                ranked = {row["chunk_id"]: {"id": row["id"], "title": row["title"], "content": row["content"], "status": row["status"], "score": 1.0 / (1.0 + abs(float(row["rank"])))} for row in rows}
                try:
                    query_vector = embedding_vectors([query])[0]
                    scope = "WHERE f.knowledge_id=?" if knowledge_id else ""; vector_rows = connection.execute(f"SELECT c.id AS chunk_id,f.id,f.name AS title,c.content,c.embedding,f.status FROM file_chunks c JOIN files f ON f.id=c.file_id {scope}", ([knowledge_id] if knowledge_id else [])).fetchall()
                    semantic = sorted(((cosine_similarity(query_vector, json.loads(row["embedding"])), row) for row in vector_rows if row["embedding"]), key=lambda item: item[0], reverse=True)[:10]
                    for similarity, row in semantic:
                        item = ranked.setdefault(row["chunk_id"], {"id": row["id"], "title": row["title"], "content": row["content"], "status": row["status"], "score": 0.0}); item["score"] = item["score"] * 0.45 + max(0.0, similarity) * 0.55
                except Exception:
                    pass
                documents, seen_files = [], set()
                for item in sorted(ranked.values(), key=lambda value: value["score"], reverse=True):
                    if item["id"] in seen_files: continue
                    seen_files.add(item["id"]); documents.append(item)
                    if len(documents) >= max(1, min(int(data.get("limit", 10)), 10)): break
                json_response(self, 200, {"documents": documents})
            elif self.path == "/api/web-search":
                query = str(data.get("query", "")).strip()
                if not query: raise ValueError("搜索内容不能为空")
                site = re.sub(r"[^a-zA-Z0-9.\-]", "", str(data.get("site", "")))
                search_query = f"site:{site} {query}" if site else query
                documents = web_search(search_query, max(1, min(int(data.get("limit", 6)), 10)))
                if site: documents = [item for item in documents if site.lower() in urllib.parse.urlparse(item.get("url", "")).netloc.lower()]
                record_usage(connection, self.identity()[0], "", "web_search", request_started); json_response(self, 200, {"documents": documents, "query": search_query})
            elif self.path == "/api/web-pages/read":
                page = read_public_web_page(data.get("url", ""))
                record_usage(connection, self.identity()[0], "", "web_page_read", request_started); json_response(self, 200, {"page": page})
            elif self.path == "/api/chat/stream":
                if not setting("api_key") and not connection.execute("SELECT 1 FROM provider_connections WHERE enabled=1 AND api_key<>'' LIMIT 1").fetchone():
                    raise ValueError("请先在模型设置中配置 API Key")
                request_messages = data.get("messages", []) if isinstance(data.get("messages", []), list) else []
                request_history = data.get("history", []) if isinstance(data.get("history", []), list) else []
                question = str(data.get("question", "")).strip(); selected_model = None
                if request_messages:
                    last_message = request_messages[-1] if isinstance(request_messages[-1], dict) else {}
                    if str(last_message.get("role", "")).strip().lower() == "user":
                        question = message_text(last_message.get("content", "")) or question
                    request_history = request_messages[:-1]
                chat_id = str(data.get("chat_id", "")).strip()
                history = conversation_messages(connection, self.identity()[0], chat_id, request_history)
                if not question and history:
                    question = message_text(history[-1].get("content", ""))
                if not question:
                    raise ValueError("问题不能为空")
                if data.get("model_id"):
                    selected_model = connection.execute("SELECT * FROM models WHERE id=?", (str(data.get("model_id")),)).fetchone()
                    if selected_model and not self.can_access_model(selected_model): raise ValueError("无权使用该模型")
                if selected_model:
                    try:
                        if not data.get("skill_ids"): data["skill_ids"] = json.loads(selected_model["skill_ids"] or "[]")
                        if not data.get("tool_ids"): data["tool_ids"] = json.loads(selected_model["tool_ids"] or "[]")
                    except (TypeError, json.JSONDecodeError):
                        pass
                filters, actions = model_runtime_config(selected_model)
                question = apply_input_filters(question, filters)
                documents = data.get("documents", []) if isinstance(data.get("documents", []), list) else []
                documents = documents + file_documents(connection, data.get("file_ids", []))
                context = "\n\n".join(f"[{index + 1}] {item.get('title', '')}\n{item.get('content', '')}" for index, item in enumerate(documents[:8]) if isinstance(item, dict))
                messages = []
                memories = connection.execute("SELECT content FROM ai_memories WHERE user_id=? AND enabled=1 ORDER BY updated_at DESC LIMIT 20", (self.identity()[0],)).fetchall()
                if memories: messages.append({"role": "system", "content": "用户授权的长期记忆：\n" + "\n".join(f"- {row['content']}" for row in memories)[:12000]})
                if selected_model and selected_model["system_prompt"]: messages.append({"role": "system", "content": selected_model["system_prompt"]})
                if actions: messages.append({"role": "system", "content": "模型动作要求：\n" + action_instructions(actions)})
                skill_ids = [str(item) for item in data.get("skill_ids", [])] if isinstance(data.get("skill_ids", []), list) else []
                if skill_ids:
                    placeholders = ",".join("?" for _ in skill_ids); skills = connection.execute(f"SELECT name,content FROM skills WHERE id IN ({placeholders})", skill_ids).fetchall()
                    if skills: messages.append({"role": "system", "content": "\n\n".join(f"# Skill: {skill['name']}\n{skill['content']}" for skill in skills)})
                tool_output = run_tools(connection, [str(item) for item in data.get("tool_ids", [])] if isinstance(data.get("tool_ids", []), list) else [], question)
                if tool_output: messages.append({"role": "system", "content": "以下是已执行工具的可信结果：\n\n" + tool_output})
                if context: messages.append({"role": "system", "content": "请优先根据以下资料回答，并标注来源编号。\n\n" + context[:20000]})
                messages.extend(history)
                image_urls = [str(item) for item in data.get("image_urls", []) if str(item).startswith(("data:image/", "https://", "http://"))][:4] if isinstance(data.get("image_urls", []), list) else []
                user_content = [{"type": "text", "text": question}] + [{"type": "image_url", "image_url": {"url": url}} for url in image_urls] if image_urls else question
                messages.append({"role": "user", "content": user_content})
                payload = {"model": selected_model["base_model"] if selected_model else setting("model", "gpt-4.1-mini"), "messages": messages, "stream": True}
                if selected_model: payload.update({"temperature": selected_model["temperature"], "top_p": selected_model["top_p"], "max_tokens": selected_model["max_tokens"]})
                provider = connection.execute("SELECT * FROM provider_connections WHERE id=? AND enabled=1", (selected_model["connection_id"],)).fetchone() if selected_model and selected_model["connection_id"] else None
                provider_type = provider["provider_type"] if provider and "provider_type" in provider.keys() else "openai"
                provider_id = str(provider["provider_id"] if provider and "provider_id" in provider.keys() else "custom").lower()
                provider_base = provider_api_base(provider["base_url"] if provider else setting("base_url", "https://api.openai.com/v1"), provider_type)
                endpoint = provider_base + ("/api/chat" if provider_type == "ollama" else "/chat/completions")
                api_key = provider["api_key"] if provider else setting("api_key")
                if provider_type == "ollama": payload = {"model": payload["model"], "messages": payload["messages"], "stream": True}
                headers = {"Content-Type": "application/json", "User-Agent": "RuoShopAdmin/1.0"};
                if api_key: headers["Authorization"] = "Bearer " + api_key
                request = urllib.request.Request(endpoint, data=json.dumps(payload).encode(), headers=headers, method="POST")

                def non_stream_answer():
                    fallback_payload = dict(payload)
                    fallback_payload["stream"] = False
                    fallback_request = urllib.request.Request(endpoint, data=json.dumps(fallback_payload).encode(), headers=headers, method="POST")
                    with urllib.request.urlopen(fallback_request, timeout=90) as fallback_response:
                        fallback_result = json.loads(fallback_response.read().decode("utf-8"))
                    if provider_type == "ollama":
                        return fallback_result.get("message", {}).get("content", "")
                    return fallback_result.get("choices", [{}])[0].get("message", {}).get("content", "")

                def write_single_answer(answer):
                    self.send_response(200); self.send_header("Content-Type", "application/x-ndjson; charset=utf-8"); self.send_header("Cache-Control", "no-cache"); self.end_headers()
                    if answer:
                        self.wfile.write((json.dumps({"content": answer}, ensure_ascii=False) + "\n").encode("utf-8")); self.wfile.flush()

                selected_tool_ids = [str(item) for item in data.get("tool_ids", [])] if isinstance(data.get("tool_ids", []), list) else []
                specs, tool_map = tool_specs(connection, selected_tool_ids)
                if specs and provider_type != "ollama":
                    try:
                        answer, _usage = complete_with_tools(endpoint, headers, payload, specs, tool_map)
                        record_usage(connection, self.identity()[0], selected_model["id"] if selected_model else "", "chat_tools", request_started, usage=_usage, cost=usage_cost(selected_model, _usage))
                        write_single_answer(answer)
                        return
                    except urllib.error.HTTPError:
                        # Some OpenAI-compatible gateways reject the tools
                        # fields. The injected-context path below remains the
                        # compatibility fallback for those providers.
                        pass

                used_responses = False
                try:
                    response = urllib.request.urlopen(request, timeout=90)
                except urllib.error.HTTPError as upstream_error:
                    try:
                        answer = non_stream_answer()
                    except urllib.error.HTTPError:
                        answer = ""
                    if answer:
                        write_single_answer(answer)
                        return
                    if provider_type != "ollama" and provider_id != "google" and upstream_error.code in (400, 404, 405, 422, 500, 502, 503):
                        fallback_endpoint = provider_base + "/responses"
                        fallback_payload = {"model": payload["model"], "input": payload["messages"], "stream": True}
                        fallback_request = urllib.request.Request(fallback_endpoint, data=json.dumps(fallback_payload).encode(), headers=headers, method="POST")
                        response = urllib.request.urlopen(fallback_request, timeout=90)
                        used_responses = True
                    else:
                        raise

                stream_meta = {"usage": {}}

                def upstream_chunks(upstream_response):
                    for raw_line in upstream_response:
                        line = raw_line.decode("utf-8", errors="replace").strip()
                        if line.startswith("data:"):
                            value = line[5:].strip()
                        elif line.startswith("{") or line.startswith("["):
                            # Ollama and a few gateways stream newline-delimited JSON.
                            value = line
                        else:
                            continue
                        if value == "[DONE]": break
                        try:
                            event = json.loads(value)
                            usage = upstream_event_usage(event)
                            if usage:
                                stream_meta["usage"] = usage
                            chunk = upstream_event_value(event)
                        except (json.JSONDecodeError, IndexError, TypeError):
                            chunk = ""
                        if chunk:
                            yield chunk

                chunks = upstream_chunks(response)
                first_chunk = next(chunks, None)
                if first_chunk is None and provider_type != "ollama" and provider_id != "google" and not used_responses:
                    response.close()
                    try:
                        answer = non_stream_answer()
                    except urllib.error.HTTPError:
                        answer = ""
                    if answer:
                        write_single_answer(answer)
                        return
                    fallback_endpoint = provider_base + "/responses"
                    fallback_payload = {"model": payload["model"], "input": payload["messages"], "stream": True}
                    fallback_request = urllib.request.Request(fallback_endpoint, data=json.dumps(fallback_payload).encode(), headers=headers, method="POST")
                    response = urllib.request.urlopen(fallback_request, timeout=90)
                    chunks = upstream_chunks(response)
                    first_chunk = next(chunks, None)

                with response:
                    self.send_response(200); self.send_header("Content-Type", "application/x-ndjson; charset=utf-8"); self.send_header("Cache-Control", "no-cache"); self.end_headers()
                    if first_chunk:
                        self.wfile.write((json.dumps({"content": first_chunk}, ensure_ascii=False) + "\n").encode("utf-8")); self.wfile.flush()
                    for chunk in chunks:
                        self.wfile.write((json.dumps({"content": chunk}, ensure_ascii=False) + "\n").encode("utf-8")); self.wfile.flush()
                record_usage(connection, self.identity()[0], selected_model["id"] if selected_model else "", "chat_stream", request_started, usage=stream_meta["usage"], cost=usage_cost(selected_model, stream_meta["usage"]))
            elif self.path == "/api/chat":
                if not setting("api_key") and not connection.execute("SELECT 1 FROM provider_connections WHERE enabled=1 AND api_key<>'' LIMIT 1").fetchone():
                    raise ValueError("请先在模型设置中配置 API Key")
                request_messages = data.get("messages", []) if isinstance(data.get("messages", []), list) else []
                request_history = data.get("history", []) if isinstance(data.get("history", []), list) else []
                question = str(data.get("question", "")).strip()
                if request_messages:
                    last_message = request_messages[-1] if isinstance(request_messages[-1], dict) else {}
                    if str(last_message.get("role", "")).strip().lower() == "user":
                        question = message_text(last_message.get("content", "")) or question
                    request_history = request_messages[:-1]
                history = conversation_messages(connection, self.identity()[0], str(data.get("chat_id", "")).strip(), request_history)
                if not question and history:
                    question = message_text(history[-1].get("content", ""))
                if not question:
                    raise ValueError("问题不能为空")
                selected_model = None
                if data.get("model_id"):
                    selected_model = connection.execute("SELECT * FROM models WHERE id=?", (str(data.get("model_id")),)).fetchone()
                    if selected_model and not self.can_access_model(selected_model): raise ValueError("无权使用该模型")
                if selected_model:
                    try:
                        if not data.get("skill_ids"): data["skill_ids"] = json.loads(selected_model["skill_ids"] or "[]")
                        if not data.get("tool_ids"): data["tool_ids"] = json.loads(selected_model["tool_ids"] or "[]")
                    except (TypeError, json.JSONDecodeError):
                        pass
                filters, actions = model_runtime_config(selected_model)
                question = apply_input_filters(question, filters)
                documents = data.get("documents", []) if isinstance(data.get("documents", []), list) else []
                documents = documents + file_documents(connection, data.get("file_ids", []))
                context = "\n\n".join(f"[{index + 1}] {item.get('title', '')}\n{item.get('content', '')}" for index, item in enumerate(documents[:8]) if isinstance(item, dict))
                messages = []
                memories = connection.execute("SELECT content FROM ai_memories WHERE user_id=? AND enabled=1 ORDER BY updated_at DESC LIMIT 20", (self.identity()[0],)).fetchall()
                if memories: messages.append({"role": "system", "content": "用户授权的长期记忆：\n" + "\n".join(f"- {row['content']}" for row in memories)[:12000]})
                if selected_model and selected_model["system_prompt"]:
                    messages.append({"role": "system", "content": selected_model["system_prompt"]})
                if actions:
                    messages.append({"role": "system", "content": "模型动作要求：\n" + action_instructions(actions)})
                skill_ids = [str(item) for item in data.get("skill_ids", [])] if isinstance(data.get("skill_ids", []), list) else []
                if skill_ids:
                    placeholders = ",".join("?" for _ in skill_ids); skills = connection.execute(f"SELECT name,content FROM skills WHERE id IN ({placeholders})", skill_ids).fetchall()
                    if skills: messages.append({"role": "system", "content": "\n\n".join(f"# Skill: {skill['name']}\n{skill['content']}" for skill in skills)})
                tool_output = run_tools(connection, [str(item) for item in data.get("tool_ids", [])] if isinstance(data.get("tool_ids", []), list) else [], question)
                if tool_output: messages.append({"role": "system", "content": "以下是已执行工具的可信结果：\n\n" + tool_output})
                if context:
                    messages.append({"role": "system", "content": "请优先根据以下资料回答，并在引用时标注来源编号。\n\n" + context[:20000]})
                messages.extend(history)
                image_urls = [str(item) for item in data.get("image_urls", []) if str(item).startswith(("data:image/", "https://", "http://"))][:4] if isinstance(data.get("image_urls", []), list) else []
                user_content = [{"type": "text", "text": question}] + [{"type": "image_url", "image_url": {"url": url}} for url in image_urls] if image_urls else question
                messages.append({"role": "user", "content": user_content})
                payload = {"model": selected_model["base_model"] if selected_model else setting("model", "gpt-4.1-mini"), "messages": messages}
                if selected_model: payload.update({"temperature": selected_model["temperature"], "top_p": selected_model["top_p"], "max_tokens": selected_model["max_tokens"]})
                provider = connection.execute("SELECT * FROM provider_connections WHERE id=? AND enabled=1", (selected_model["connection_id"],)).fetchone() if selected_model and selected_model["connection_id"] else None
                provider_type = provider["provider_type"] if provider and "provider_type" in provider.keys() else "openai"
                provider_base = provider_api_base(provider["base_url"] if provider else setting("base_url", "https://api.openai.com/v1"), provider_type)
                endpoint = provider_base + ("/api/chat" if provider_type == "ollama" else "/chat/completions")
                api_key = provider["api_key"] if provider else setting("api_key")
                if provider_type == "ollama": payload = {"model": payload["model"], "messages": payload["messages"], "stream": False}
                headers = {"Content-Type": "application/json", "User-Agent": "RuoShopAdmin/1.0"};
                if api_key: headers["Authorization"] = "Bearer " + api_key
                request = urllib.request.Request(endpoint, data=json.dumps(payload).encode(), headers=headers, method="POST")
                with urllib.request.urlopen(request, timeout=90) as response:
                    result = json.loads(response.read().decode())
                answer = result.get("message", {}).get("content", "") if provider_type == "ollama" else result.get("choices", [{}])[0].get("message", {}).get("content", "")
                record_usage(connection, self.identity()[0], selected_model["id"] if selected_model else "", "chat", request_started, usage=result.get("usage", {}), cost=usage_cost(selected_model, result.get("usage", {}))); json_response(self, 200, {"answer": answer})
            elif self.path == "/api/files/generate":
                title = str(data.get("title", "AI 输出")).strip()[:120] or "AI 输出"; content = str(data.get("content", ""))
                if not content.strip(): raise ValueError("文件内容不能为空")
                file_format = str(data.get("format", "docx")).lower(); raw, mime = generate_document(title, content, file_format)
                record_usage(connection, self.identity()[0], "", "file_generate", request_started); json_response(self, 200, {"filename": re.sub(r'[\\/:*?\"<>|]+', '_', title) + "." + file_format, "mime": mime, "data": base64.b64encode(raw).decode("ascii")})
            elif self.path == "/api/audio/transcriptions":
                raw = base64.b64decode(str(data.get("data", "")), validate=True); filename = Path(str(data.get("filename", "audio.webm"))).name
                if not raw: raise ValueError("音频不能为空")
                model_id = str(data.get("model_id", "")); model = connection.execute("SELECT * FROM models WHERE id=? AND enabled=1 AND model_type='audio'", (model_id,)).fetchone() if model_id else connection.execute("SELECT * FROM models WHERE enabled=1 AND model_type='audio' ORDER BY is_default DESC,pinned DESC,name LIMIT 1").fetchone()
                if not model: raise ValueError("没有可用的语音转写模型")
                provider = connection.execute("SELECT * FROM provider_connections WHERE id=? AND enabled=1", (model["connection_id"],)).fetchone()
                if not provider: raise ValueError("语音模型没有绑定启用的账号连接")
                body, boundary = multipart_file({"model": model["base_model"]}, "file", filename, raw, audio_content_type(filename))
                endpoint = provider_api_base(provider["base_url"], provider["provider_type"]) + "/audio/transcriptions"
                request = urllib.request.Request(endpoint, data=body, headers={"Content-Type": "multipart/form-data; boundary=" + boundary, "Authorization": "Bearer " + provider["api_key"], "User-Agent": "RuoShopAdmin/1.0"}, method="POST")
                with urllib.request.urlopen(request, timeout=180) as response: result = json.loads(response.read().decode("utf-8"))
                record_usage(connection, self.identity()[0], model["id"], "audio_transcription", request_started, usage=result.get("usage", {})); json_response(self, 200, {"text": str(result.get("text", ""))})
            elif self.path == "/api/audio/speech":
                text_value = str(data.get("text", "")).strip()
                if not text_value: raise ValueError("朗读内容不能为空")
                if len(text_value) > 4096: raise ValueError("朗读内容过长，请分段播放")
                model_id = str(data.get("model_id", "")); model = connection.execute("SELECT * FROM models WHERE id=? AND enabled=1 AND model_type='audio'", (model_id,)).fetchone() if model_id else connection.execute("SELECT * FROM models WHERE enabled=1 AND model_type='audio' ORDER BY is_default DESC,pinned DESC,name LIMIT 1").fetchone()
                if not model: raise ValueError("没有可用的语音模型")
                provider = connection.execute("SELECT * FROM provider_connections WHERE id=? AND enabled=1", (model["connection_id"],)).fetchone()
                if not provider: raise ValueError("语音模型没有绑定启用的账号连接")
                endpoint = provider_api_base(provider["base_url"], provider["provider_type"]) + "/audio/speech"
                payload = {"model": model["base_model"], "input": text_value, "voice": str(data.get("voice", "alloy")), "response_format": "mp3"}
                request = urllib.request.Request(endpoint, data=json.dumps(payload).encode(), headers={"Content-Type": "application/json", "Authorization": "Bearer " + provider["api_key"], "User-Agent": "RuoShopAdmin/1.0"}, method="POST")
                with urllib.request.urlopen(request, timeout=180) as response: audio = response.read()
                record_usage(connection, self.identity()[0], model["id"], "audio_speech", request_started); json_response(self, 200, {"mime": "audio/mpeg", "data": base64.b64encode(audio).decode("ascii")})
            elif self.path == "/api/images/generations":
                if not setting("api_key") and not connection.execute("SELECT 1 FROM provider_connections WHERE enabled=1 AND api_key<>'' LIMIT 1").fetchone():
                    raise ValueError("请先在模型设置中配置 API Key")
                prompt_text = str(data.get("prompt", "")).strip()
                if not prompt_text: raise ValueError("图片描述不能为空")
                selected_model = connection.execute("SELECT * FROM models WHERE id=?", (str(data.get("model_id", "")),)).fetchone() if data.get("model_id") else None
                if selected_model and not self.can_access_model(selected_model): raise ValueError("无权使用该模型")
                if not selected_model or ("model_type" in selected_model.keys() and selected_model["model_type"] != "image"):
                    selected_model = connection.execute("SELECT * FROM models WHERE enabled=1 AND model_type='image' ORDER BY is_default DESC,pinned DESC,name LIMIT 1").fetchone()
                if not selected_model: raise ValueError("没有可用的图片模型，请先同步或创建 gpt-image、DALL-E、Flux 等图片模型")
                provider = connection.execute("SELECT * FROM provider_connections WHERE id=? AND enabled=1", (selected_model["connection_id"],)).fetchone() if selected_model and selected_model["connection_id"] else None
                provider_type = provider["provider_type"] if provider and "provider_type" in provider.keys() else "openai"
                if provider_type == "ollama": raise ValueError("Ollama 当前不支持图片生成接口")
                provider_base = provider_api_base(provider["base_url"] if provider else setting("base_url", "https://api.openai.com/v1"), provider_type)
                api_key = provider["api_key"] if provider else setting("api_key")
                payload = {"prompt": prompt_text, "model": selected_model["base_model"] if selected_model else setting("model", "gpt-image-1"), "size": str(data.get("size", "1024x1024")), "n": 1}
                reference_images = data.get("image_urls", []) if isinstance(data.get("image_urls", []), list) else []
                reference_image = str(reference_images[0]) if reference_images else ""
                if reference_image:
                    match = re.fullmatch(r"data:(image/(?:png|jpeg|webp));base64,(.+)", reference_image, re.DOTALL)
                    if not match: raise ValueError("参考图格式无效，仅支持 PNG、JPG 或 WebP")
                    try: image_raw = base64.b64decode(match.group(2), validate=True)
                    except Exception as error: raise ValueError("参考图数据损坏") from error
                    if len(image_raw) > 15_000_000: raise ValueError("参考图不能超过 15MB")
                    suffix = "jpg" if match.group(1) == "image/jpeg" else match.group(1).split("/", 1)[1]
                    request_body, boundary = multipart_file(payload, "image", "reference." + suffix, image_raw, match.group(1))
                    endpoint = provider_base + "/images/edits"
                    content_type = "multipart/form-data; boundary=" + boundary
                else:
                    endpoint = provider_base + "/images/generations"
                    request_body = json.dumps(payload).encode()
                    content_type = "application/json"
                request = urllib.request.Request(endpoint, data=request_body, headers={"Content-Type": content_type, "Authorization": "Bearer " + api_key, "User-Agent": "RuoShopAdmin/1.0"}, method="POST")
                try:
                    with urllib.request.urlopen(request, timeout=180) as response:
                        result = json.loads(response.read().decode("utf-8"))
                except urllib.error.HTTPError as error:
                    detail = error.read().decode("utf-8", errors="replace")[:2000]
                    action = "图生图" if reference_image else "图片生成"
                    raise ValueError(f"{action}供应商返回 HTTP {error.code}: {detail or error.reason}") from error
                item = (result.get("data") or [{}])[0]
                image_url = item.get("url", "")
                if not image_url and item.get("b64_json"):
                    image_url = "data:image/png;base64," + item["b64_json"]
                if not image_url: raise ValueError("图片服务未返回图片地址")
                operation = "image_edit" if reference_image else "image_generation"
                record_usage(connection, self.identity()[0], selected_model["id"], operation, request_started, usage=result.get("usage", {})); json_response(self, 200, {"url": image_url, "revised_prompt": item.get("revised_prompt", ""), "mode": "edit" if reference_image else "generation"})
            elif self.path == "/api/test":
                if not setting("api_key"):
                    raise ValueError("请先配置 API Key")
                endpoint = provider_api_base(setting("base_url", "https://api.openai.com/v1")) + "/chat/completions"
                payload = {"model": setting("model", "gpt-4.1-mini"), "messages": [{"role": "user", "content": "请只回复：连接成功"}], "max_tokens": 20}
                request = urllib.request.Request(endpoint, data=json.dumps(payload).encode(), headers={"Content-Type": "application/json", "Authorization": "Bearer " + setting("api_key"), "User-Agent": "RuoShopAdmin/1.0"}, method="POST")
                with urllib.request.urlopen(request, timeout=30) as response:
                    result = json.loads(response.read().decode())
                json_response(self, 200, {"ok": True, "answer": result.get("choices", [{}])[0].get("message", {}).get("content", "连接成功")})
            else:
                json_response(self, 404, {"error": "接口不存在"})
        except PermissionError as error:
            json_response(self, 401, {"error": str(error)})
        except RateLimitError as error:
            json_response(self, 429, {"error": str(error), "retry_after": 60})
        except Exception as error:
            json_response(self, 400, {"error": str(error)})
        finally:
            if connection is not None:
                connection.close()


if __name__ == "__main__":
    db().close(); ThreadingHTTPServer((HOST, PORT), Handler).serve_forever()
